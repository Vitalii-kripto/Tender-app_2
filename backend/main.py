import uvicorn
import asyncio
import sys
import os
import logging
import tempfile
import re
import sqlite3
import threading
import uuid
from functools import partial
from dotenv import load_dotenv

# Загружаем переменные окружения в самом начале
load_dotenv()

from backend.logging_setup import setup_logging
# Инициализируем логирование максимально рано
setup_logging("LegalAI")

from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor

from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Body, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
import shutil
from starlette.concurrency import run_in_threadpool

if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

from .database import engine, Base, get_db, SessionLocal
from .models import TenderModel, ProductModel
from .services.eis_service import EisService, Notice, mark_seen
from backend.config import DOCUMENTS_ROOT
from .services.parser import GidroizolParser
from .services.document_service import DocumentService
from .services.ai_service import AiService
from .services.legal_analysis_service import LegalAnalysisService
from .services.goods_extraction_service import GoodsExtractionService
from .services.tender_document_inventory import collect_visible_tender_files, wait_for_tender_input_files
from .services.batch_analysis import analyze_tenders_batch_job
from .services.job_service import job_service
from .services.analog_search_job_service import analog_search_job_service
from pydantic import BaseModel

class FrontendLog(BaseModel):
    level: str = "info"
    message: str
    context: Dict[str, Any] = {}

def clean_markdown(text):
    """Удаляет markdown-артефакты из текста"""
    if not text: return ""
    # Удаляем жирный/курсив
    text = re.sub(r'\*\*|\*|__|_', '', text)
    # Удаляем заголовки
    text = re.sub(r'#+\s+', '', text)
    # Удаляем горизонтальные линии
    text = re.sub(r'^-{3,}\s*$', '', text, flags=re.MULTILINE)
    # Удаляем лишние пустые строки
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def parse_markdown_table(text):
    """Парсит markdown-таблицу в список списков"""
    if not text or '|' not in text: return []
    lines = text.strip().split('\n')
    table_rows = []
    for line in lines:
        if '|' in line:
            # Пропускаем разделители |---|---|
            if re.match(r'^[\s|:\-\+]+$', line.strip()):
                continue
            # Разбиваем по |
            parts = [p.strip() for p in line.split('|')]
            # Убираем пустые элементы по краям, если строка начиналась/заканчивалась на |
            if line.strip().startswith('|'):
                parts = parts[1:]
            if line.strip().endswith('|'):
                parts = parts[:-1]
            
            if parts and any(p for p in parts):
                table_rows.append([clean_markdown(p) for p in parts])
    return table_rows


_active_docs_downloads_lock = threading.Lock()
_active_docs_downloads: set[str] = set()


def _run_tender_doc_download(tender_id: str, notice: Notice) -> None:
    try:
        eis_service.redownload_tender_documents(notice, clear_dir=True)
    finally:
        with _active_docs_downloads_lock:
            _active_docs_downloads.discard(tender_id)


def _schedule_tender_doc_download(
    background_tasks: BackgroundTasks,
    tender_id: str,
    notice: Notice,
) -> bool:
    with _active_docs_downloads_lock:
        if tender_id in _active_docs_downloads:
            logger.info(
                "Skipped duplicate document download scheduling for tender %s: download already in progress",
                tender_id,
            )
            return False
        _active_docs_downloads.add(tender_id)

    background_tasks.add_task(_run_tender_doc_download, tender_id, notice)
    logger.info(f"Scheduled document download for tender {tender_id}")
    return True

def parse_markdown_list(text):
    """Парсит markdown-список в список строк"""
    if not text: return []
    lines = text.strip().split('\n')
    list_items = []
    for line in lines:
        # Ищем маркеры списка: -, *, 1., 1)
        match = re.match(r'^\s*(?:[\-\*\+]|\d+[\.\)])\s+(.*)', line)
        if match:
            list_items.append(clean_markdown(match.group(1)))
        elif line.strip() and not re.match(r'^[\s|:\-\+]+$', line.strip()) and '|' not in line:
            # Если это просто строка текста, тоже берем ее как элемент, если она не пустая
            list_items.append(clean_markdown(line))
    return list_items


def parse_price_to_float(raw_price: Any) -> float:
    if raw_price is None or raw_price == "":
        return 0.0
    if isinstance(raw_price, (int, float)):
        return float(raw_price)

    text = str(raw_price).replace("\xa0", " ").strip()
    cleaned = re.sub(r"[^\d,.\-]", "", text)

    if not cleaned:
        return 0.0

    # Наиболее частый кейс для RU-формата
    if cleaned.count(",") == 1 and cleaned.count(".") == 0:
        cleaned = cleaned.replace(",", ".")
    elif cleaned.count(",") > 0 and cleaned.count(".") > 0:
        # Если и точка, и запятая — считаем последнюю запятую десятичным разделителем
        if cleaned.rfind(",") > cleaned.rfind("."):
            cleaned = cleaned.replace(".", "").replace(",", ".")
        else:
            cleaned = cleaned.replace(",", "")
    else:
        cleaned = cleaned.replace(",", "")

    try:
        return float(cleaned)
    except Exception:
        return 0.0


def normalize_law_type(
    *,
    reg: str = "",
    ntype: str = "",
    url: str = "",
    docs_url: str = "",
    explicit: Any = None,
) -> str:
    explicit_text = str(explicit or "").strip()
    if explicit_text in {"44-ФЗ", "223-ФЗ", "Коммерч."}:
        return explicit_text

    reg = str(reg or "")
    haystack = " ".join([
        reg,
        str(ntype or ""),
        str(url or ""),
        str(docs_url or ""),
    ]).lower()

    if reg.startswith("223-") or "notice223" in haystack:
        return "223-ФЗ"

    if "коммер" in haystack or "commercial" in haystack:
        return "Коммерч."

    return "44-ФЗ"


def build_docs_url_from_payload(payload: Dict[str, Any]) -> str:
    docs_url = str(payload.get("docs_url") or "").strip()
    if docs_url:
        return docs_url

    reg = str(payload.get("id") or payload.get("eis_number") or "").strip()
    ntype = str(payload.get("ntype") or "").strip()

    if not reg:
        return ""

    if reg.startswith("223-"):
        notice_id = reg.replace("223-", "", 1).strip()
        if notice_id:
            return f"https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?noticeInfoId={notice_id}"
        return ""

    if ntype:
        return f"https://zakupki.gov.ru/epz/order/notice/{ntype}/view/documents.html?regNumber={reg}"

    return ""


def tender_has_local_files(tender_id: str) -> bool:
    tender_dir = os.path.join(DOCUMENTS_ROOT, tender_id)
    if not os.path.isdir(tender_dir):
        return False
    return any(os.path.isfile(os.path.join(tender_dir, name)) for name in os.listdir(tender_dir))


def tender_model_to_payload(tender: TenderModel) -> Dict[str, Any]:
    return {
        "id": tender.id,
        "eis_number": tender.id,
        "title": tender.title or "",
        "description": tender.description or "",
        "initial_price": tender.initial_price or 0,
        "deadline": tender.deadline or "",
        "status": tender.status or "Found",
        "risk_level": tender.risk_level or "Low",
        "region": tender.region or "РФ",
        "law_type": normalize_law_type(
            reg=tender.id,
            ntype=tender.ntype or "",
            url=tender.url or "",
            docs_url=tender.docs_url or "",
            explicit=tender.law_type,
        ),
        "url": tender.url or "",
        "docs_url": tender.docs_url or "",
        "search_url": tender.search_url or "",
        "keyword": tender.keyword or "",
        "ntype": tender.ntype or "",
        "customer_name": tender.customer_name or "",
        "customer_inn": tender.customer_inn or "",
        "customer_location": tender.customer_location or "",
        "selected_for_matching": bool(tender.selected_for_matching),
    }


def _normalize_manual_text_value(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\xa0", " ")).strip()


def _extract_first_regex(text: str, patterns: List[str], flags: int = re.IGNORECASE) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags)
        if match:
            return _normalize_manual_text_value(match.group(1))
    return ""


def _fallback_extract_tender_details_from_text(text: str) -> Dict[str, Any]:
    normalized_text = text or ""
    compact_text = _normalize_manual_text_value(normalized_text)

    eis_numbers = re.findall(r"\b\d{11,19}\b", normalized_text)
    eis_number = max(eis_numbers, key=len) if eis_numbers else ""

    law_type = ""
    lowered = normalized_text.lower()
    if "223-фз" in lowered or "223 фз" in lowered:
        law_type = "223-ФЗ"
    elif "44-фз" in lowered or "44 фз" in lowered:
        law_type = "44-ФЗ"
    elif "коммер" in lowered or "запрос предложений" in lowered:
        law_type = "Коммерч."

    title = _extract_first_regex(
        normalized_text,
        [
            r"(?:наименование\s+(?:закупки|объекта закупки)|предмет\s+(?:закупки|договора))[:\s]+(.{10,200})",
            r"(?:объект\s+закупки)[:\s]+(.{10,200})",
        ],
    )
    if not title and compact_text:
        for line in normalized_text.splitlines():
            candidate = _normalize_manual_text_value(line)
            if len(candidate) < 12:
                continue
            if re.search(r"(утверждаю|документац|извещен|протокол|техническое задание)", candidate, re.IGNORECASE):
                continue
            title = candidate[:200]
            break

    deadline = _extract_first_regex(
        normalized_text,
        [
            r"(?:дата\s+окончания\s+подачи\s+заявок|окончание\s+срока\s+подачи\s+заявок|срок\s+подачи\s+заявок)[:\s]+(\d{2}\.\d{2}\.\d{4}(?:\s+\d{2}:\d{2})?)",
            r"(?:заявок\s+до)[:\s]+(\d{2}\.\d{2}\.\d{4}(?:\s+\d{2}:\d{2})?)",
        ],
    )
    publish_date = _extract_first_regex(
        normalized_text,
        [
            r"(?:дата\s+(?:размещени[яе]|публикации)|размещено)[:\s]+(\d{2}\.\d{2}\.\d{4})",
        ],
    )

    customer_name = _extract_first_regex(
        normalized_text,
        [
            r"(?:заказчик|наименование\s+заказчика)[:\s]+(.{5,200})",
        ],
    )
    customer_inn = _extract_first_regex(
        normalized_text,
        [
            r"(?:инн(?:\s+заказчика)?)[:\s]+(\d{10,12})",
        ],
    )
    customer_location = _extract_first_regex(
        normalized_text,
        [
            r"(?:место\s+нахождения\s+заказчика|адрес\s+заказчика)[:\s]+(.{5,200})",
        ],
    )

    price_match = re.search(
        r"(?:нмцк|начальная\s*\(максимальная\)\s*цена(?:\s+договора|\s+контракта)?|цена\s+договора)[:\s]*([\d\s\xa0,\.]+)",
        normalized_text,
        re.IGNORECASE,
    )
    initial_price = parse_price_to_float(price_match.group(1)) if price_match else 0.0

    description = ""
    if compact_text:
        description = compact_text[:2000]

    return {
        "title": title,
        "initial_price": initial_price,
        "deadline": deadline,
        "publish_date": publish_date,
        "eis_number": eis_number,
        "description": description,
        "law_type": law_type,
        "customer_name": customer_name,
        "customer_inn": customer_inn,
        "customer_location": customer_location,
    }


def _merge_manual_tender_details(
    ai_details: Optional[Dict[str, Any]],
    fallback_details: Dict[str, Any],
) -> Dict[str, Any]:
    merged = dict(fallback_details)
    for key, value in (ai_details or {}).items():
        if value is None:
            continue
        if isinstance(value, str):
            cleaned = _normalize_manual_text_value(value)
            if cleaned:
                merged[key] = cleaned
            continue
        if key == "initial_price":
            parsed = parse_price_to_float(value)
            if parsed > 0:
                merged[key] = parsed
            continue
        merged[key] = value

    if not merged.get("law_type"):
        merged["law_type"] = fallback_details.get("law_type") or "Коммерч."

    return merged


def build_notice_from_payload(payload: Dict[str, Any]) -> Notice:
    reg = str(payload.get("id") or payload.get("eis_number") or "").strip()
    docs_url = build_docs_url_from_payload(payload)

    return Notice(
        reg=reg,
        ntype=str(payload.get("ntype") or "").strip(),
        keyword=str(payload.get("keyword") or "").strip(),
        search_url=str(payload.get("search_url") or "").strip(),
        href=str(payload.get("url") or "").strip(),
        docs_url=docs_url,
        title=str(payload.get("title") or "").strip(),
        object_info=str(payload.get("description") or "").strip(),
        initial_price=str(payload.get("initial_price") or "").strip(),
        application_deadline=str(payload.get("deadline") or "").strip(),
        customer_name=str(payload.get("customer_name") or "").strip(),
        customer_inn=str(payload.get("customer_inn") or "").strip(),
        customer_location=str(payload.get("customer_location") or "").strip(),
    )


def save_or_update_tender_record(
    *,
    db: Session,
    tender_payload: Dict[str, Any],
    background_tasks: BackgroundTasks,
    force_download: bool = False,
) -> TenderModel:
    tender_id = str(tender_payload.get("id") or "").strip()
    if not tender_id:
        raise HTTPException(status_code=400, detail="У тендера отсутствует id")

    title = str(tender_payload.get("title") or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail=f"У тендера {tender_id} отсутствует title")

    docs_url = build_docs_url_from_payload(tender_payload)
    law_type = normalize_law_type(
        reg=tender_id,
        ntype=str(tender_payload.get("ntype") or ""),
        url=str(tender_payload.get("url") or ""),
        docs_url=docs_url,
        explicit=tender_payload.get("law_type"),
    )
    parsed_price = parse_price_to_float(tender_payload.get("initial_price", 0))
    selected_for_matching = bool(tender_payload.get("selected_for_matching", False))

    existing = db.query(TenderModel).filter(TenderModel.id == tender_id).first()

    if existing:
        previous_docs_url = existing.docs_url or ""

        existing.title = title
        existing.description = str(tender_payload.get("description") or "")
        existing.initial_price = parsed_price
        existing.deadline = str(tender_payload.get("deadline") or "")
        existing.status = str(tender_payload.get("status") or existing.status or "Found")
        existing.risk_level = str(tender_payload.get("risk_level") or existing.risk_level or "Low")
        existing.region = str(tender_payload.get("region") or existing.region or "РФ")
        existing.law_type = law_type
        existing.url = str(tender_payload.get("url") or "")
        existing.docs_url = docs_url
        existing.search_url = str(tender_payload.get("search_url") or "")
        existing.keyword = str(tender_payload.get("keyword") or "")
        existing.ntype = str(tender_payload.get("ntype") or "")
        existing.customer_name = str(tender_payload.get("customer_name") or "")
        existing.customer_inn = str(tender_payload.get("customer_inn") or "")
        existing.customer_location = str(tender_payload.get("customer_location") or "")
        existing.selected_for_matching = selected_for_matching

        db.flush()

        need_download = (
            force_download
            or (bool(docs_url) and not tender_has_local_files(existing.id))
            or (bool(docs_url) and docs_url != previous_docs_url)
        )

        model = existing
        logger.info(f"Updated existing tender: {tender_id}")
    else:
        model = TenderModel(
            id=tender_id,
            title=title,
            description=str(tender_payload.get("description") or ""),
            initial_price=parsed_price,
            deadline=str(tender_payload.get("deadline") or ""),
            status=str(tender_payload.get("status") or "Found"),
            risk_level=str(tender_payload.get("risk_level") or "Low"),
            region=str(tender_payload.get("region") or "РФ"),
            law_type=law_type,
            url=str(tender_payload.get("url") or ""),
            docs_url=docs_url,
            search_url=str(tender_payload.get("search_url") or ""),
            keyword=str(tender_payload.get("keyword") or ""),
            ntype=str(tender_payload.get("ntype") or ""),
            customer_name=str(tender_payload.get("customer_name") or ""),
            customer_inn=str(tender_payload.get("customer_inn") or ""),
            customer_location=str(tender_payload.get("customer_location") or ""),
            selected_for_matching=selected_for_matching,
        )
        db.add(model)
        db.flush()

        need_download = bool(docs_url)
        logger.info(f"Created new tender: {tender_id}")

    if need_download and docs_url:
        notice = build_notice_from_payload(
            {
                **tender_payload,
                "id": tender_id,
                "title": title,
                "docs_url": docs_url,
                "law_type": law_type,
            }
        )
        _schedule_tender_doc_download(background_tasks, tender_id, notice)

    return model

from backend.logging_setup import setup_logging, get_logger
logger = setup_logging("LegalAI")
frontend_logger = get_logger("Frontend")

# --- SETUP ---
def migrate_db():
    """Простейшая миграция для добавления недостающих колонок в SQLite"""
    from .database import DB_PATH
    
    logger.info(f"Checking schema for {DB_PATH}...")
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Список колонок, которые могли быть добавлены позже
    # (column_name, table_name, column_type)
    new_columns = [
        ("docs_url", "tenders", "TEXT"),
        ("search_url", "tenders", "TEXT"),
        ("keyword", "tenders", "TEXT"),
        ("ntype", "tenders", "TEXT"),
        ("customer_name", "tenders", "TEXT"),
        ("customer_inn", "tenders", "TEXT"),
        ("customer_location", "tenders", "TEXT"),
        ("local_file_path", "tenders", "TEXT"),
        ("extracted_text", "tenders", "TEXT"),
        ("selected_for_matching", "tenders", "BOOLEAN DEFAULT 0"),
        ("created_at", "tenders", "DATETIME"),
        ("description", "products", "TEXT"),
        ("updated_at", "products", "DATETIME"),
        ("vendor", "products", "TEXT DEFAULT 'gidroizol'"),
        ("site_product_id", "products", "TEXT"),
        ("source_url", "products", "TEXT"),
        ("city_id", "products", "INTEGER DEFAULT 1"),
        ("category_leaf", "products", "TEXT"),
        ("normalized_category", "products", "TEXT"),
        ("searchable_for_analogs", "products", "BOOLEAN DEFAULT 1"),
        ("price_wholesale", "products", "FLOAT"),
        ("price_special", "products", "FLOAT"),
        ("price_currency", "products", "TEXT DEFAULT 'RUB'"),
        ("price_unit", "products", "TEXT"),
        ("availability_status", "products", "TEXT DEFAULT 'unknown'"),
        ("specs_text", "products", "TEXT"),
        ("meta_description", "products", "TEXT"),
        ("quality_score", "products", "INTEGER DEFAULT 0"),
        ("content_hash", "products", "TEXT"),
        ("parse_version", "products", "TEXT"),
        ("first_seen_at", "products", "DATETIME"),
        ("last_seen_at", "products", "DATETIME"),
        ("scraped_at", "products", "DATETIME"),
        ("is_active", "products", "BOOLEAN DEFAULT 1"),
    ]
    
    for col_name, table_name, col_type in new_columns:
        try:
            cursor.execute(f"ALTER TABLE {table_name} ADD COLUMN {col_name} {col_type}")
            logger.info(f"Added column {col_name} to table {table_name}")
        except sqlite3.OperationalError as e:
            if "duplicate column name" in str(e).lower():
                pass
            else:
                logger.critical(f"Migration error for {table_name}.{col_name}: {e}")

    ddl_statements = [
        """
        CREATE TABLE IF NOT EXISTS product_analog_index (
            product_id INTEGER PRIMARY KEY,
            normalized_title TEXT NOT NULL,
            brand TEXT,
            series TEXT,
            product_family TEXT,
            material_group TEXT NOT NULL,
            material_subgroup TEXT,
            application_scope TEXT,
            base_material TEXT,
            binder_type TEXT,
            top_surface TEXT,
            bottom_surface TEXT,
            thickness_mm FLOAT,
            mass_kg_m2 FLOAT,
            density_kg_m3 FLOAT,
            roll_length_m FLOAT,
            roll_width_m FLOAT,
            roll_area_m2 FLOAT,
            package_weight_kg FLOAT,
            flexibility_temp_c FLOAT,
            heat_resistance_c FLOAT,
            color TEXT,
            standard_code TEXT,
            extracted_attrs_json JSON DEFAULT '{}',
            search_text TEXT NOT NULL,
            analog_group_key TEXT,
            exact_model_key TEXT,
            updated_at DATETIME
        )
        """,
        """
        CREATE TABLE IF NOT EXISTS product_aliases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            product_id INTEGER NOT NULL,
            alias TEXT NOT NULL,
            alias_normalized TEXT NOT NULL,
            alias_type TEXT NOT NULL,
            CONSTRAINT uq_product_aliases_product_norm UNIQUE (product_id, alias_normalized)
        )
        """,
        "CREATE UNIQUE INDEX IF NOT EXISTS idx_products_url_unique ON products(url)",
        "CREATE INDEX IF NOT EXISTS idx_products_vendor ON products(vendor)",
        "CREATE INDEX IF NOT EXISTS idx_products_category_leaf ON products(category_leaf)",
        "CREATE INDEX IF NOT EXISTS idx_products_searchable_quality ON products(searchable_for_analogs, quality_score)",
        "CREATE INDEX IF NOT EXISTS idx_products_active ON products(is_active, last_seen_at)",
        "CREATE INDEX IF NOT EXISTS idx_products_material_type ON products(material_type)",
        "CREATE INDEX IF NOT EXISTS idx_analog_group ON product_analog_index(material_group, material_subgroup)",
        "CREATE INDEX IF NOT EXISTS idx_analog_family ON product_analog_index(product_family)",
        "CREATE INDEX IF NOT EXISTS idx_analog_brand_series ON product_analog_index(brand, series)",
        "CREATE INDEX IF NOT EXISTS idx_analog_base ON product_analog_index(base_material)",
        "CREATE INDEX IF NOT EXISTS idx_analog_thickness ON product_analog_index(thickness_mm)",
        "CREATE INDEX IF NOT EXISTS idx_alias_norm ON product_aliases(alias_normalized)",
    ]
    for ddl in ddl_statements:
        try:
            cursor.execute(ddl)
        except sqlite3.OperationalError as e:
            logger.critical(f"Migration DDL error: {e} | SQL={ddl}")

    try:
        cursor.execute(
            """
            CREATE VIRTUAL TABLE IF NOT EXISTS product_search_fts
            USING fts5(
                product_id UNINDEXED,
                normalized_title,
                search_text,
                specs_text,
                description,
                aliases,
                tokenize = 'unicode61 remove_diacritics 2'
            )
            """
        )
    except sqlite3.OperationalError as e:
        logger.warning(f"FTS5 initialization skipped: {e}")

    conn.commit()
    conn.close()

try:
    logger.info("Initializing Database...")
    Base.metadata.create_all(bind=engine)
    migrate_db()
    logger.info("Database initialized successfully.")
except Exception as e:
    logger.critical(f"Database initialization failed: {e}", exc_info=True)

app = FastAPI(title="TenderSmart Gidroizol API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Разрешить все для локальной разработки
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Services
eis_service = None
parser_service = None
doc_service = None
ai_service = None
legal_analysis_service = None
analog_service = None
_analog_search_refinement_lock = None
_analog_search_refinement_retry_tasks = set()
_ANALOG_SEARCH_RETRY_DELAYS_SECONDS = (30, 60, 180)
_ANALOG_SEARCH_QUOTA_MAX_RETRIES = 2

def db_session_factory():
    return SessionLocal()


async def run_blocking(func, *args, **kwargs):
    return await run_in_threadpool(partial(func, *args, **kwargs))


def _get_analog_search_refinement_lock():
    global _analog_search_refinement_lock
    if _analog_search_refinement_lock is None:
        _analog_search_refinement_lock = asyncio.Lock()
    return _analog_search_refinement_lock


def _track_analog_search_refinement_retry_task(task: asyncio.Task) -> None:
    _analog_search_refinement_retry_tasks.add(task)

    def _cleanup(completed_task: asyncio.Task) -> None:
        _analog_search_refinement_retry_tasks.discard(completed_task)
        try:
            completed_task.result()
        except asyncio.CancelledError:
            pass
        except Exception:
            logger.error("Analog search delayed retry task failed", exc_info=True)

    task.add_done_callback(_cleanup)


def _build_analog_search_warnings(query: str) -> List[str]:
    warnings: List[str] = []
    if len(query) < 10 or (not bool(re.search(r'\d|[a-zA-Z]', query)) and len(query.split()) < 3):
        warnings.append(
            "Поисковый запрос выглядит слишком общим "
            "(вероятно из-за низкого качества извлечения ТЗ). Данные поиска могут быть нерелевантными."
        )
    return warnings


def _apply_analog_search_mode(result: Dict[str, Any], mode: str) -> Dict[str, Any]:
    payload = dict(result or {})
    payload["local_results"] = list(payload.get("local_results") or [])
    payload["ai_results"] = list(payload.get("ai_results") or [])

    if mode == "ai":
        payload["local_results"] = []
    elif mode == "local":
        payload["ai_results"] = []

    payload["total"] = len(payload["local_results"]) + len(payload["ai_results"])
    payload["mode"] = mode
    payload["ai_error"] = payload.get("ai_error", "") or ""
    return payload


def _merge_analog_ai_error(payload: Dict[str, Any]) -> Dict[str, Any]:
    result = dict(payload or {})
    ai_error_parts = [
        part for part in [
            result.get("ai_error", ""),
            result.get("validation_error", ""),
        ]
        if part
    ]
    result["ai_error"] = " | ".join(ai_error_parts)
    return result


def _extract_retryable_analog_ai_reason(error_text: str) -> str:
    normalized_error = str(error_text or "").upper()
    if not normalized_error:
        return ""
    if any(marker in normalized_error for marker in ["QUOTA_EXHAUSTED", "RESOURCE_EXHAUSTED", "429"]):
        return "QUOTA_EXHAUSTED"
    if any(marker in normalized_error for marker in ["SERVICE_UNAVAILABLE", "503"]):
        return "SERVICE_UNAVAILABLE"
    return ""


def _get_analog_search_retry_limit(retry_reason: str = "") -> int:
    if str(retry_reason or "").upper() == "QUOTA_EXHAUSTED":
        return _ANALOG_SEARCH_QUOTA_MAX_RETRIES
    return len(_ANALOG_SEARCH_RETRY_DELAYS_SECONDS)


def _get_analog_search_retry_delay(retry_count: int, retry_reason: str = "") -> int | None:
    if retry_count < 0:
        retry_count = 0
    retry_limit = _get_analog_search_retry_limit(retry_reason)
    if retry_count >= retry_limit:
        return None
    return _ANALOG_SEARCH_RETRY_DELAYS_SECONDS[retry_count]


def _append_analog_ai_error(payload: Dict[str, Any] | None, message: str) -> Dict[str, Any]:
    result = dict(payload or {})
    error_parts: list[str] = []
    for part in [result.get("ai_error", ""), str(message or "").strip()]:
        part = str(part or "").strip()
        if part and part not in error_parts:
            error_parts.append(part)
    result["ai_error"] = " | ".join(error_parts)
    return result


def _build_exhausted_analog_search_result(
    payload: Dict[str, Any] | None,
    *,
    reason: str,
    retry_count: int,
) -> Dict[str, Any]:
    if str(reason or "").upper() == "QUOTA_EXHAUSTED":
        exhausted_message = (
            f"Internet-поиск недоступен по квоте после {retry_count} попыток: {reason}"
        )
    else:
        exhausted_message = f"Отложенные повторы исчерпаны после {retry_count} попыток: {reason}"
    return _append_analog_ai_error(payload, exhausted_message)


async def _delayed_analog_search_refinement_retry(
    *,
    job_id: str,
    query: str,
    requirements: str | None,
    mode: str,
    max_results: int,
    delay_seconds: int,
) -> None:
    logger.info(
        "Analog search delayed retry scheduled | job_id=%s | delay=%ss | query=%s",
        job_id,
        delay_seconds,
        query[:120],
    )
    await asyncio.sleep(max(1, int(delay_seconds)))

    job = analog_search_job_service.get_job(job_id)
    if not job:
        logger.info("Analog search delayed retry skipped | job_id=%s | reason=job_missing", job_id)
        return
    if job.get("status") != "waiting_retry":
        logger.info(
            "Analog search delayed retry skipped | job_id=%s | status=%s",
            job_id,
            job.get("status"),
        )
        return

    await _run_analog_search_refinement_job(
        job_id=job_id,
        query=query,
        requirements=requirements,
        mode=mode,
        max_results=max_results,
    )


def _schedule_analog_search_retry(
    *,
    job_id: str,
    query: str,
    requirements: str | None,
    mode: str,
    max_results: int,
    result: Dict[str, Any],
    retry_reason: str,
    delay_override_seconds: int | None = None,
) -> bool:
    job = analog_search_job_service.get_job(job_id)
    if not job:
        return False

    retry_count = int(job.get("retry_count") or 0)
    delay_seconds = max(1, int(delay_override_seconds or 0)) if delay_override_seconds else None
    if delay_seconds is None:
        delay_seconds = _get_analog_search_retry_delay(retry_count, retry_reason)
    if delay_seconds is None:
        analog_search_job_service.complete_job(
            job_id,
            _build_exhausted_analog_search_result(
                result,
                reason=retry_reason,
                retry_count=retry_count,
            ),
        )
        logger.warning(
            "Analog search refinement retries exhausted | job_id=%s | retries=%s | reason=%s",
            job_id,
            retry_count,
            retry_reason,
        )
        return True

    scheduled_job = analog_search_job_service.schedule_retry(
        job_id,
        delay_seconds=delay_seconds,
        reason=retry_reason,
        result=result,
    )
    if not scheduled_job:
        return False

    retry_task = asyncio.create_task(
        _delayed_analog_search_refinement_retry(
            job_id=job_id,
            query=query,
            requirements=requirements,
            mode=mode,
            max_results=max_results,
            delay_seconds=delay_seconds,
        )
    )
    _track_analog_search_refinement_retry_task(retry_task)
    return True


async def _run_analog_search_refinement_job(
    *,
    job_id: str,
    query: str,
    requirements: str | None,
    mode: str,
    max_results: int,
) -> None:
    current_job = analog_search_job_service.get_job(job_id) or {}
    current_context = analog_search_job_service.get_job_context(job_id) or {}
    current_retry_count = int(current_job.get("retry_count") or 0)
    current_result = current_job.get("result")
    current_base_result = current_context.get("base_result") or current_result
    lock = _get_analog_search_refinement_lock()
    if lock.locked():
        logger.info(
            "Analog search refinement queued | job_id=%s | query=%s",
            job_id,
            query[:120],
        )

    async with lock:
        stage = "Ищу аналоги в интернете"
        if current_retry_count > 0:
            current_retry_reason = _extract_retryable_analog_ai_reason(current_job.get("error", ""))
            retry_limit = _get_analog_search_retry_limit(current_retry_reason)
            stage = (
                f"Повторная попытка internet/AI-уточнения "
                f"({current_retry_count}/{retry_limit})"
            )

        analog_search_job_service.start_job(job_id, stage=stage)
        try:
            raw_result = await analog_service.refine_search_result(
                base_result=current_base_result or {},
                query=query,
                requirements=requirements,
                use_ai=(mode in ["ai", "both"]),
                limit=max_results,
                retry_internet_only=current_retry_count > 0,
            )
            analog_search_job_service.update_job_context(
                job_id,
                {"base_result": raw_result},
            )
            final_result = _merge_analog_ai_error(_apply_analog_search_mode(raw_result, mode))
            retry_reason = _extract_retryable_analog_ai_reason(raw_result.get("ai_error", ""))
            if retry_reason:
                retry_delay_override = None
                if retry_reason == "QUOTA_EXHAUSTED":
                    retry_delay_override = max(
                        int(analog_service._remaining_ai_block_seconds("grounded_search") or 0),
                        int(analog_service._remaining_query_ai_block_seconds(
                            query,
                            requirements or "",
                            scope="grounded_search",
                        ) or 0),
                    )
                if _schedule_analog_search_retry(
                    job_id=job_id,
                    query=query,
                    requirements=requirements,
                    mode=mode,
                    max_results=max_results,
                    result=final_result,
                    retry_reason=retry_reason,
                    delay_override_seconds=retry_delay_override,
                ):
                    return
            analog_search_job_service.complete_job(job_id, final_result)
        except Exception as e:
            retry_reason = _extract_retryable_analog_ai_reason(str(e))
            if retry_reason:
                fallback_result = _append_analog_ai_error(current_result, retry_reason)
                retry_delay_override = None
                if retry_reason == "QUOTA_EXHAUSTED":
                    retry_delay_override = max(
                        int(analog_service._remaining_ai_block_seconds("grounded_search") or 0),
                        int(analog_service._remaining_query_ai_block_seconds(
                            query,
                            requirements or "",
                            scope="grounded_search",
                        ) or 0),
                    )
                if _schedule_analog_search_retry(
                    job_id=job_id,
                    query=query,
                    requirements=requirements,
                    mode=mode,
                    max_results=max_results,
                    result=fallback_result,
                    retry_reason=retry_reason,
                    delay_override_seconds=retry_delay_override,
                ):
                    return
            logger.error(
                "Analog search refinement job failed | job_id=%s | query=%s | error=%s",
                job_id,
                query[:120],
                e,
                exc_info=True,
            )
            analog_search_job_service.fail_job(job_id, str(e))

logger.info("Initializing Services...")

try:
    eis_service = EisService()
    logger.info("[OK] EisService initialized.")
except Exception as e:
    logger.error(f"[FAILED] EisService initialization error: {e}")

try:
    doc_service = DocumentService()
    logger.info("[OK] DocumentService initialized.")
except Exception as e:
    logger.error(f"[FAILED] DocumentService initialization error: {e}")

try:
    ai_service = AiService()
    if not ai_service.has_available_clients():
        logger.critical("[WARNING] AI Service is UNAVAILABLE. Backend will start in degraded mode.")
    else:
        logger.info(
            "[OK] AI Service initialized. Startup connectivity check skipped to preserve quota. "
            "Route selection will happen on the first real AI request."
        )

    try:
        legal_analysis_service = LegalAnalysisService(ai_service)
        logger.info("[OK] LegalAnalysisService initialized.")
    except Exception as e:
        logger.error(f"[FAILED] LegalAnalysisService initialization error: {e}")

    try:
        goods_extraction_service = GoodsExtractionService(ai_service)
        logger.info("[OK] GoodsExtractionService initialized.")
    except Exception as e:
        logger.error(f"[FAILED] GoodsExtractionService initialization error: {e}")

except Exception as e:
    logger.error(f"[FAILED] AiService initialization error: {e}")

try:
    parser_service = GidroizolParser()
    logger.info("[OK] GidroizolParser initialized.")
except Exception as e:
    logger.error(f"[FAILED] GidroizolParser initialization error: {e}")

try:
    from backend.services.analog_service import AnalogService

    if not ai_service:
        raise RuntimeError("AiService is not initialized, AnalogService cannot be created.")

    analog_service = AnalogService(
        ai_service=ai_service,
        db_session_factory=db_session_factory
    )
    logger.info("[OK] AnalogService initialized.")
except Exception as e:
    logger.error(f"[FAILED] AnalogService initialization error: {e}")

logger.info("Service initialization phase completed.")

def build_notice_from_tender_model(tender: TenderModel) -> Notice:
    return Notice(
        reg=tender.id,
        ntype=tender.ntype or "",
        keyword=tender.keyword or "",
        search_url=tender.search_url or "",
        href=tender.url or "",
        docs_url=tender.docs_url or "",
        title=tender.title or "",
        object_info=tender.description or "",
        initial_price=str(tender.initial_price or ""),
        application_deadline=tender.deadline or "",
    )

def _matching_norm(value: str) -> str:
    value = (value or "").lower().replace("\xa0", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def _text_has_product_signals(text: str) -> bool:
    low = _matching_norm(text)
    if not low:
        return False

    score = 0

    material_terms = [
        "гидроизол", "техноэласт", "унифлекс", "линокром", "биполь", "филизол",
        "рубитэкс", "эластобит", "стеклоэласт", "стеклоизол", "гидростеклоизол",
        "изоэласт", "мостослой", "тэксослой", "мастика", "праймер", "мембрана",
        "герметик", "лента", "геотекстиль", "геомембрана", "шпонка", "пароизоляция"
    ]
    tech_terms = [
        "характерист", "основа", "толщин", "масса", "гибкост", "теплостойк",
        "маркировк", "рулон", "кг", "мм", "м2", "м²", "ед. изм", "количество"
    ]

    if any(term in low for term in material_terms):
        score += 3
    if any(term in low for term in tech_terms):
        score += 2
    if re.search(r"\b(тпп|ткп|хпп|хкп|эпп|экп|эмп)\b", low, flags=re.IGNORECASE):
        score += 4
    if re.search(r"\b\d+(?:[.,]\d+)?\s*(мм|м2|м²|м|кг|л|шт|рулон)\b", low, flags=re.IGNORECASE):
        score += 2

    return score >= 4


def classify_document_for_matching(filename: str, extracted_text: str) -> str:
    name = _matching_norm(filename)
    text = _matching_norm((extracted_text or "")[:12000])
    joined = f"{name} {text}".strip()

    if not joined:
        return "other"

    if any(term in joined for term in ["нмцд", "нмцк", "обоснован", "расчет нмц", "смет"]):
        return "price_calc"

    if any(term in joined for term in [
        "нац режим", "национальн", "страна происх", "реестр российск",
        "ограничени", "запрет", "преференц"
    ]):
        return "national_regime"

    if any(term in joined for term in [
        "инструкц", "заявк", "образц", "форма ", "требовани к содержанию",
        "критерии оценки", "участник закупки"
    ]):
        return "application_docs"

    if any(term in joined for term in [
        "техническое задание", "техзадание", "тз", "описание объекта закупки",
        "спецификац", "ведомость материалов", "перечень материалов",
        "товарная номенклатура", "технические характеристики"
    ]):
        return "tender_spec"

    if any(term in joined for term in ["договор", "контракт", "проект договора", "проект контракта"]):
        if _text_has_product_signals(text):
            return "contract_mixed"
        return "contract_legal"

    if any(term in joined for term in ["извещение", "notice", "информационная карта"]):
        return "notice"

    if _text_has_product_signals(joined):
        return "mixed"

    return "other"


def fragment_relevance_for_matching(fragment: str, role: str) -> int:
    low = _matching_norm(fragment)
    if not low:
        return 0

    score = 0

    positive_terms = [
        "техническое задание", "описание объекта закупки", "спецификац",
        "наименование товара", "характерист", "материал", "товар", "поставка",
        "основа", "толщин", "масса", "гибкост", "теплостойк"
    ]
    negative_terms = [
        "оплата", "штраф", "пеня", "неустойк", "расторжен", "ответственность",
        "реквизит", "участник закупки", "инструкция", "критерии оценки",
        "нмцд", "нмцк", "обоснован", "нац режим", "реестр", "банковская гарантия",
        "обеспечение исполнения", "срок оплаты", "порядок расчетов"
    ]

    if role in {"tender_spec", "contract_mixed", "mixed"}:
        score += 3

    if any(term in low for term in positive_terms):
        score += 4

    if re.search(r"\b(тпп|ткп|хпп|хкп|эпп|экп|эмп)\b", low, flags=re.IGNORECASE):
        score += 4

    if re.search(r"\b\d+(?:[.,]\d+)?\s*(мм|м2|м²|м|кг|л|шт|рулон)\b", low, flags=re.IGNORECASE):
        score += 2

    if _text_has_product_signals(low):
        score += 3

    negative_hits = sum(1 for term in negative_terms if term in low)
    score -= negative_hits * 2

    if re.search(r"\.(docx?|pdf|xlsx?|xls)\b", low, flags=re.IGNORECASE):
        score -= 3
    if re.search(r"\bстр\.?\s*\d+\b", low, flags=re.IGNORECASE):
        score -= 2

    if len(low) < 25:
        score -= 2

    return score


def extract_requirement_fragments_from_document_data(
    doc_data: Dict[str, Any],
    role: str,
    max_chars_per_file: int = 24000
) -> List[str]:
    filename = doc_data.get("filename", "Unknown")
    pages = doc_data.get("pages", []) or []
    fragments_with_score: List[tuple[int, str]] = []

    def _push_fragment(page_num: Any, raw_text: str):
        if not raw_text:
            return

        raw_text = str(raw_text).replace("\xa0", " ")
        raw_text = re.sub(r"\n{3,}", "\n\n", raw_text).strip()
        if not raw_text:
            return

        for block in re.split(r"\n\s*\n+", raw_text):
            block = block.strip()
            if not block:
                continue

            score = fragment_relevance_for_matching(block, role)
            threshold = 4 if role in {"tender_spec", "contract_mixed", "mixed"} else 7
            if score < threshold:
                continue

            cleaned = block[:3500].strip()
            if not cleaned:
                continue

            header = f"[FILE: {filename} | PAGE: {page_num} | ROLE: {role} | SCORE: {score}]"
            fragments_with_score.append((score, header + "\n" + cleaned))

    for page in pages:
        page_num = page.get("page_num", "")
        page_text = page.get("text", "") or ""
        tables = page.get("tables", []) or []

        _push_fragment(page_num, page_text)

        for table_text in tables:
            _push_fragment(page_num, table_text)

    fragments_with_score.sort(key=lambda item: item[0], reverse=True)

    selected: List[str] = []
    total_chars = 0
    for _, fragment in fragments_with_score:
        if total_chars + len(fragment) > max_chars_per_file:
            break
        selected.append(fragment)
        total_chars += len(fragment)

    return selected

# --- DEGRADED MODE CHECKS ---
def check_eis_service():
    if not eis_service:
        raise HTTPException(status_code=503, detail="EIS Service is not initialized.")

def check_parser_service():
    if not parser_service:
        raise HTTPException(status_code=503, detail="Parser Service is not initialized.")

def check_doc_service():
    if not doc_service:
        raise HTTPException(status_code=503, detail="Document Service is not initialized.")

def check_ai_service():
    if not ai_service or not ai_service.client:
        raise HTTPException(status_code=503, detail="AI Service is not initialized or unavailable.")

def check_legal_service():
    if not legal_analysis_service:
        raise HTTPException(status_code=503, detail="Legal Analysis Service is not initialized.")

def check_analog_service():
    if not analog_service:
        raise HTTPException(status_code=503, detail="Analog Service is not initialized.")

# --- ENDPOINTS ---

@app.post("/api/frontend-log")
async def post_frontend_log(log_data: FrontendLog):
    """Принимает frontend-логи и пишет их в общий tendersmart.txt"""
    f_logger = logging.getLogger("Frontend")

    msg = f"[FRONTEND][{log_data.level.upper()}] {log_data.message}"
    if log_data.context:
        msg += f" | Context: {log_data.context}"

    level = log_data.level.lower().strip()
    if level == "error":
        f_logger.error(msg)
    elif level == "warning":
        f_logger.warning(msg)
    else:
        f_logger.info(msg)

    return {"status": "ok"}

@app.get("/")
def read_root():
    logger.info("Health check endpoint hit.")
    return {"status": "online", "system": "TenderSmart PRO Backend"}

# --- CRM ENDPOINTS (Database Sync) ---

@app.get("/api/crm/tenders")
def get_crm_tenders(db: Session = Depends(get_db)):
    """Получить все тендеры из базы"""
    logger.info("Fetching all CRM tenders.")
    return db.query(TenderModel).all()

@app.post("/api/crm/tenders")
def add_update_tender(
    background_tasks: BackgroundTasks,
    tender: dict = Body(...),
    db: Session = Depends(get_db),
):
    """Добавить или обновить тендер в CRM"""
    tender_id = str(tender.get("id") or "").strip()
    logger.info(f"Add/Update tender request: {tender_id}")

    try:
        save_or_update_tender_record(
            db=db,
            tender_payload=tender,
            background_tasks=background_tasks,
            force_download=False,
        )
        db.commit()
        return {"status": "success"}
    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Error saving tender: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/crm/tenders/{tender_id}/requeue-docs")
def requeue_tender_docs(tender_id: str, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    """Повторно запустить скачивание документов для тендера"""
    tender = db.query(TenderModel).filter(TenderModel.id == tender_id).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Tender not found")
        
    if not tender.docs_url:
        raise HTTPException(status_code=400, detail="Tender has no docs_url")
        
    notice = Notice(
        reg=tender.id,
        ntype=tender.ntype or "",
        keyword=tender.keyword or "",
        search_url=tender.search_url or "",
        title=tender.title or "",
        href=tender.url or "",
        object_info=tender.description or "",
        initial_price=str(tender.initial_price or ""),
        application_deadline=tender.deadline or "",
        docs_url=tender.docs_url
    )
    _schedule_tender_doc_download(background_tasks, tender_id, notice)
    return {"status": "requeued"}

@app.delete("/api/crm/tenders/{tender_id}")
def delete_tender(tender_id: str, db: Session = Depends(get_db)):
    """Удалить тендер из базы"""
    logger.info(f"Deleting tender: {tender_id}")
    tender = db.query(TenderModel).filter(TenderModel.id == tender_id).first()
    if tender:
        db.delete(tender)
        db.commit()
    return {"status": "deleted"}

# --- SEARCH & PARSING ---

@app.post("/api/search-tenders/cancel")
def cancel_search():
    """Отменить текущий поиск"""
    logger.info("Cancel search request received")
    eis_service.cancel_search()
    return {"status": "cancelled"}

@app.get("/api/search-tenders")
def search_tenders_endpoint(
    query: str,
    fz44: bool = True,
    fz223: bool = True,
    only_application_stage: bool = True,
    publish_days_back: int = 30,
    _ = Depends(check_eis_service),
):
    """Поиск через Playwright"""
    import uuid

    search_id = uuid.uuid4().hex[:8]
    logger.info(f"[SEARCH_API_START] [SearchID:{search_id}] Search request received: {query}")
    try:
        notices = eis_service.search_tenders(
            query=query,
            fz44=fz44,
            fz223=fz223,
            only_application_stage=only_application_stage,
            publish_days_back=publish_days_back,
            search_id=search_id,
        )
        logger.info(
            f"[SEARCH_SERVICE_RESPONSE] [SearchID:{search_id}] "
            f"eis_service returned {len(notices)} items before customer enrichment"
        )
    except RuntimeError as e:
        logger.error(f"[SEARCH_API_ERROR] [SearchID:{search_id}] Search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    except Exception as e:
        logger.error(f"[SEARCH_API_ERROR] [SearchID:{search_id}] Search failed with unexpected error: {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

    # Convert Notice dataclass to dict for JSON response
    result = []
    for n in notices:
        try:
            eis_service.enrich_notice_customer_info(n)
        except Exception as exc:
            logger.warning(
                "Failed to enrich customer info for tender %s: %s",
                n.reg,
                exc,
            )

        docs_url = build_docs_url_from_payload({
            "id": n.reg,
            "ntype": n.ntype,
            "docs_url": n.docs_url
        })
        law_type = normalize_law_type(
            reg=n.reg,
            ntype=n.ntype,
            url=n.href,
            docs_url=docs_url
        )
        parsed_price = parse_price_to_float(n.initial_price)

        result.append({
            "id": n.reg,
            "eis_number": n.reg,
            "title": n.title,
            "description": n.object_info,
            "initial_price": parsed_price,
            "initial_price_text": str(n.initial_price),
            "deadline": n.application_deadline,
            "status": "Found",
            "risk_level": "Low",
            "region": "РФ",
            "law_type": law_type,
            "url": n.href,
            "docs_url": docs_url,
            "search_url": n.search_url,
            "keyword": n.keyword,
            "ntype": n.ntype,
            "customer_name": n.customer_name,
            "customer_inn": n.customer_inn,
            "customer_location": n.customer_location,
        })

    logger.info(
        f"[SEARCH_API_RESPONSE] [SearchID:{search_id}] "
        f"returning {len(result)} items after customer enrichment"
    )
    return result

@app.post("/api/search-tenders/process")
def process_tenders(
    background_tasks: BackgroundTasks,
    tenders: list = Body(...),
    db: Session = Depends(get_db),
):
    """Массовая обработка выбранных тендеров"""
    if not tenders:
        return {"status": "ok", "processed": 0}

    logger.info(f"Processing {len(tenders)} selected tenders")

    for t_payload in tenders:
        try:
            save_or_update_tender_record(
                db=db,
                tender_payload=t_payload,
                background_tasks=background_tasks,
                force_download=False,
            )
        except Exception as e:
            logger.error(f"Failed to process tender in batch: {e}", exc_info=True)

    db.commit()
    return {"status": "ok", "processed": len(tenders)}


@app.post("/api/search-tenders/manual-upload")
async def upload_manual_tender_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db),
    _ = Depends(check_doc_service),
):
    """Создать тендер из вручную загруженной документации."""
    valid_files = [file for file in files if file and (file.filename or "").strip()]
    if not valid_files:
        raise HTTPException(status_code=400, detail="Не переданы файлы документации")

    temp_upload_id = f"manual_upload_{uuid.uuid4().hex[:12]}"
    tender_dir = os.path.join(DOCUMENTS_ROOT, temp_upload_id)
    os.makedirs(tender_dir, exist_ok=True)

    uploaded_paths: List[str] = []
    try:
        for file in valid_files:
            safe_name = os.path.basename(file.filename.strip())
            if not safe_name:
                continue
            file_path = os.path.join(tender_dir, safe_name)
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            uploaded_paths.append(file_path)

        if not uploaded_paths:
            raise HTTPException(status_code=400, detail="Не удалось сохранить загруженные файлы")

        try:
            from backend.services.archive_service import archive_service
            archive_service.unpack_directory(tender_dir)
        except Exception as exc:
            logger.error("Manual tender archive unpack failed: %s", exc, exc_info=True)

        all_files, has_archives, _ = await run_blocking(
            wait_for_tender_input_files,
            "",
            temp_upload_id,
            DOCUMENTS_ROOT,
            logger=logger,
            max_wait_seconds=2.0,
            poll_interval_seconds=0.5,
            settle_seconds=0.5,
        )
        if not all_files:
            all_files = [
                path for path in uploaded_paths
                if os.path.exists(path) and os.path.isfile(path)
            ]

        extracted_chunks: List[str] = []
        extracted_files_count = 0
        extraction_errors: List[str] = []

        for file_path in all_files:
            extracted = await run_blocking(
                doc_service.extract_document_data,
                file_path,
                tender_id=temp_upload_id,
            )
            if extracted.get("status") == "success" and (extracted.get("extracted_text") or "").strip():
                extracted_files_count += 1
                extracted_chunks.append(
                    f"=== FILE: {os.path.basename(file_path)} ===\n{extracted.get('extracted_text', '')}"
                )
            else:
                extraction_errors.append(
                    f"{os.path.basename(file_path)}: {extracted.get('error_message') or extracted.get('status') or 'extract_error'}"
                )

        combined_text = "\n\n".join(extracted_chunks).strip()
        if not combined_text:
            raise HTTPException(
                status_code=400,
                detail="Не удалось извлечь текст из загруженной документации",
            )

        ai_details = await run_blocking(ai_service.extract_tender_details, combined_text) if ai_service else {}
        fallback_details = _fallback_extract_tender_details_from_text(combined_text)
        merged_details = _merge_manual_tender_details(ai_details, fallback_details)

        extracted_eis_number = str(merged_details.get("eis_number") or "").strip()
        tender_id = extracted_eis_number or f"manual_{uuid.uuid4().hex[:12]}"
        title = str(merged_details.get("title") or "").strip() or f"Загруженный тендер {tender_id}"
        description = str(merged_details.get("description") or "").strip()
        publish_date = str(merged_details.get("publish_date") or "").strip()
        if publish_date:
            description = (description + f"\n\nДата публикации: {publish_date}").strip()
        if extraction_errors:
            description = (
                description
                + "\n\nПроблемы чтения файлов:\n"
                + "\n".join(f"- {message}" for message in extraction_errors[:10])
            ).strip()

        tender_payload = {
            "id": tender_id,
            "eis_number": tender_id,
            "title": title,
            "description": description,
            "initial_price": merged_details.get("initial_price") or 0,
            "deadline": str(merged_details.get("deadline") or "").strip(),
            "status": "Found",
            "risk_level": "Low",
            "region": "Ручная загрузка",
            "law_type": merged_details.get("law_type") or "Коммерч.",
            "customer_name": merged_details.get("customer_name") or "",
            "customer_inn": merged_details.get("customer_inn") or "",
            "customer_location": merged_details.get("customer_location") or "",
        }

        model = save_or_update_tender_record(
            db=db,
            tender_payload=tender_payload,
            background_tasks=background_tasks,
            force_download=False,
        )

        final_tender_dir = os.path.join(DOCUMENTS_ROOT, tender_id)
        if os.path.abspath(final_tender_dir) != os.path.abspath(tender_dir):
            if os.path.isdir(final_tender_dir):
                shutil.rmtree(final_tender_dir, ignore_errors=True)
            shutil.move(tender_dir, final_tender_dir)
        else:
            final_tender_dir = tender_dir

        model.local_file_path = final_tender_dir
        model.extracted_text = combined_text[:200000]
        db.commit()
        db.refresh(model)

        payload = tender_model_to_payload(model)
        payload["manual_upload"] = True
        payload["documents_summary"] = {
            "uploaded_files_count": len(uploaded_paths),
            "processed_files_count": extracted_files_count,
            "has_archives": has_archives,
            "files": [os.path.basename(path) for path in all_files],
        }
        return payload
    except HTTPException:
        db.rollback()
        shutil.rmtree(tender_dir, ignore_errors=True)
        raise
    except Exception as exc:
        db.rollback()
        shutil.rmtree(tender_dir, ignore_errors=True)
        logger.error("Manual tender upload failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"Ошибка обработки документации: {exc}")

@app.post("/api/search-tenders/skip")
def skip_tender(tender: dict = Body(...)):
    """Пропустить тендер (отметить как просмотренный)"""
    logger.info(f"Skipping tender: {tender.get('id')}")
    try:
        mark_seen(tender['id'])
        return {"status": "success"}
    except Exception as e:
        logger.error(f"Error skipping tender: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/products")
def get_products_endpoint(db: Session = Depends(get_db)):
    """Получение сохраненных товаров из БД без запуска парсера"""
    logger.info("Fetching products from DB.")
    products = db.query(ProductModel).all()
    result = []
    for p in products:
        result.append({
            "id": str(p.id),
            "title": p.title,
            "category": p.category,
            "material_type": p.material_type,
            "price": p.price,
            "specs": p.specs if p.specs else {},
            "url": p.url,
            "description": p.description # Added description
        })
    return result

@app.get("/api/parse-catalog")
async def parse_catalog_endpoint(db: Session = Depends(get_db), _ = Depends(check_parser_service)):
    """Запуск парсера каталога Gidroizol.ru и обновление БД"""
    logger.info("Starting catalog parser manually.")

    parse_stats = await parser_service.parse_and_save(db)
    logger.info(
        "Parser finished. Parsed %s products | inserted=%s | updated=%s | missing_url=%s | unseen_existing=%s.",
        parse_stats.get("parsed", 0),
        parse_stats.get("inserted", 0),
        parse_stats.get("updated", 0),
        parse_stats.get("missing_url", 0),
        parse_stats.get("not_seen_existing", 0),
    )
    
    # После парсинга забираем актуальные данные из БД
    products = db.query(ProductModel).all()
    
    result = []
    for p in products:
        result.append({
            "id": str(p.id),
            "title": p.title,
            "category": p.category,
            "material_type": p.material_type,
            "price": p.price,
            "specs": p.specs if p.specs else {},
            "url": p.url,
            "description": p.description # Added description
        })
    return result

def collect_tender_text_for_matching(tender: TenderModel) -> Dict[str, Any]:
    """
    Собирает источник для подбора аналогов не из целых документов,
    а только из релевантных товарно-технических фрагментов.
    """
    fragments: List[str] = []
    warnings: List[Dict[str, Any]] = []
    files_used: List[str] = []
    fragments_meta: List[Dict[str, Any]] = []

    tender_dir = os.path.join(DOCUMENTS_ROOT, tender.id)
    files_found = 0

    if os.path.isdir(tender_dir) and doc_service:
        visible_files, archive_files = collect_visible_tender_files(tender_dir, logger=logger)
        files_found = len(visible_files)
        if archive_files:
            logger.info(
                "Skipping archive files from matching context | tender_id=%s | archives=%s",
                tender.id,
                len(archive_files),
            )

        for file_path in sorted(visible_files):
            filename = os.path.relpath(file_path, tender_dir).replace("\\", "/")

            try:
                data = doc_service.extract_document_data(file_path)
            except Exception as e:
                warnings.append({
                    "filename": filename,
                    "status": "extract_error",
                    "message": str(e),
                })
                continue

            status = data.get("status", "")
            extracted_text = (data.get("extracted_text") or "").strip()

            if status != "success" or not extracted_text:
                warnings.append({
                    "filename": filename,
                    "status": status or "empty",
                    "message": data.get("error_message") or "Текст не извлечен",
                })
                continue

            role = classify_document_for_matching(filename, extracted_text)
            relevant_fragments = extract_requirement_fragments_from_document_data(data, role)

            if not relevant_fragments:
                warnings.append({
                    "filename": filename,
                    "status": "no_relevant_fragments",
                    "message": f"Документ прочитан, но релевантные товарно-технические фрагменты не найдены (role={role})",
                })
                continue

            fragments.extend(relevant_fragments)
            files_used.append(filename)
            fragments_meta.append({
                "filename": filename,
                "role": role,
                "fragments_count": len(relevant_fragments),
            })

    # CRM description использовать только как крайний fallback,
    # если у тендера нет локальных документов вообще.
    if not fragments and files_found == 0 and (tender.description or "").strip():
        fragments.append(f"[CRM_DESCRIPTION_FALLBACK]\\n{tender.description.strip()}")
        warnings.append({
            "filename": "CRM_DESCRIPTION",
            "status": "warning",
            "message": "Документы тендера отсутствуют. Использовано только описание из CRM как аварийный fallback.",
        })

    return {
        "text": "\\n\\n".join(fragments).strip(),
        "warnings": warnings,
        "files_used": files_used,
        "fragments_meta": fragments_meta,
    }

# --- AI & DOCS ENDPOINTS ---

@app.get("/api/tenders/{tender_id}/files")
def get_tender_files(tender_id: str, _ = Depends(check_doc_service)):
    """Получить список скачанных файлов для тендера"""
    logger.info(f"Fetching files for tender {tender_id}")
    tender_dir = os.path.join(DOCUMENTS_ROOT, tender_id)
    if not os.path.exists(tender_dir):
        return []
    
    try:
        from backend.services.archive_service import archive_service
        # Распаковываем старые/ручные архивы, если они есть
        archive_service.unpack_directory(tender_dir)
    except Exception as e:
        logger.error(f"Error unpacking docs in get_tender_files: {e}")

    visible_files, archive_files = collect_visible_tender_files(tender_dir, logger=logger)
    if archive_files:
        logger.info(
            "Hiding archive files from tender UI list | tender_id=%s | archives=%s",
            tender_id,
            len(archive_files),
        )
    files = []
    for filepath in sorted(visible_files):
        try:
            file_size = os.path.getsize(filepath)
        except OSError:
            continue
        files.append({
            "name": os.path.relpath(filepath, tender_dir).replace("\\", "/"),
            "size": file_size,
            "ext": os.path.splitext(filepath)[1].lower(),
        })

    return files

import shutil

@app.post("/api/tenders/{tender_id}/refresh-files")
async def refresh_tender_files(
    tender_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """Очистить кэш файлов и скачать заново"""
    tender = db.query(TenderModel).filter(TenderModel.id == tender_id).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Тендер не найден в БД")

    tender_dir = os.path.join(DOCUMENTS_ROOT, tender_id)
    if os.path.exists(tender_dir):
        logger.info("Deleting files for tender %s before refresh", tender_id)
        shutil.rmtree(tender_dir, ignore_errors=True)

    payload = tender_model_to_payload(tender)
    notice = build_notice_from_payload(payload)

    result = await run_in_threadpool(eis_service.redownload_tender_documents, notice)
    meta = (result or {}).get("meta", {}) or {}
    reason = (result or {}).get("reason", "unknown")

    files = []
    archive_files = []
    if os.path.exists(tender_dir):
        visible_files, archive_files = collect_visible_tender_files(tender_dir, logger=logger)
        for f_path in sorted(visible_files):
            try:
                file_size = os.path.getsize(f_path)
            except OSError:
                continue
            files.append({
                "name": os.path.relpath(f_path, tender_dir).replace("\\", "/"),
                "size": file_size,
                "path": f_path,
            })

    if result and result.get("ok") and files:
        logger.info("Refresh tender files completed for %s | files=%s", tender_id, len(files))
        return {
            "status": "success",
            "tender_id": tender_id,
            "files": files,
            "count": len(files),
            "has_archives": bool(archive_files),
            "archives_count": len(archive_files),
            "errors": meta.get("failed_downloads", []),
        }

    errors = meta.get("failed_downloads", [])
    if not errors:
        errors = [{
            "reason": reason,
            "message": meta.get("reason", reason),
            "docs_url": meta.get("docs_url", ""),
        }]

    logger.warning(
        "Refresh tender files partial result for %s | reason=%s | errors=%s",
        tender_id,
        reason,
        len(errors),
    )

    return {
        "status": "partial",
        "tender_id": tender_id,
        "files": files,
        "count": len(files),
        "has_archives": bool(archive_files),
        "archives_count": len(archive_files),
        "reason": reason,
        "errors": errors,
        "meta": meta,
    }

@app.post("/api/ai/analyze-tenders-batch")
async def api_analyze_tenders_batch(background_tasks: BackgroundTasks, data: dict = Body(...), _ = Depends(check_legal_service)):
    logger.info("Batch AI Analysis request received.")
    
    tender_ids = data.get('tender_ids', [])
    selected_files = data.get('selected_files', {}) # {tender_id: [filenames]}
    
    if not tender_ids:
        raise HTTPException(status_code=400, detail="No tender IDs provided")
    
    job_id = job_service.create_job(tender_ids)
    background_tasks.add_task(analyze_tenders_batch_job, job_id, tender_ids, doc_service, legal_analysis_service, selected_files)
    
    return {"job_id": job_id}

@app.get("/api/ai/jobs/{job_id}")
async def get_job_status(job_id: str):
    job = job_service.get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job

@app.post("/api/ai/extract-details")
async def api_extract_details(data: dict = Body(...), _ = Depends(check_ai_service)):
    """Извлечение данных о тендере из текста"""
    logger.info("AI Extract Details request.")
    text = data.get('text', '')
    return await run_blocking(ai_service.extract_tender_details, text)

@app.post("/api/ai/extract-products")
async def api_extract_products(data: dict = Body(...), _ = Depends(check_ai_service)):
    """Извлечение списка товаров из сметы/КП"""
    logger.info("AI Extract Products request.")
    text = data.get('text', '')
    return await run_blocking(ai_service.extract_products_from_text, text)

def _unique_non_empty(values: List[Any]) -> List[str]:
    result: List[str] = []
    seen = set()

    for value in values or []:
        text = str(value or "").strip()
        if not text or text in seen:
            continue
        seen.add(text)
        result.append(text)

    return result


def _normalize_extraction_warnings(
    warnings: Any,
    *,
    tender_id: str = "",
    filename: str = "",
    status: str = "",
) -> List[Dict[str, Any]]:
    normalized: List[Dict[str, Any]] = []

    if not isinstance(warnings, list):
        return normalized

    for warning in warnings:
        if isinstance(warning, dict):
            message = str(warning.get("message") or warning.get("detail") or "").strip()
            if not message:
                continue
            normalized.append({
                "tender_id": warning.get("tender_id") or tender_id or None,
                "filename": warning.get("filename") or filename or None,
                "status": warning.get("status") or status or None,
                "message": message,
            })
            continue

        message = str(warning or "").strip()
        if not message:
            continue
        normalized.append({
            "tender_id": tender_id or None,
            "filename": filename or None,
            "status": status or None,
            "message": message,
        })

    return normalized


def _build_requirement_items(
    *,
    positions: List[Dict[str, Any]],
    tender_id: str,
    tender_title: str,
    source: str,
    source_label: str,
) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []

    for index, pos in enumerate(positions or []):
        name = pos.get("position_name_normalized") or pos.get("position_name_raw") or ""
        raw_chars = pos.get("characteristics") or []
        formatted_chars: List[str] = []
        structured_chars: List[Dict[str, str]] = []
        key_chars: List[str] = []

        for char in raw_chars:
            if isinstance(char, dict):
                char_name = str(char.get("name") or "").strip()
                char_value = str(char.get("value") or "").strip()
                if char_name or char_value:
                    structured_chars.append({
                        "name": char_name,
                        "value": char_value,
                    })
                formatted = (
                    f"{char_name}: {char_value}"
                    if char_name and char_value
                    else char_value or char_name
                )
                if formatted:
                    formatted_chars.append(formatted)
                char_name_lc = char_name.lower()
                if char_value and any(
                    key in char_name_lc
                    for key in ["толщина", "масса", "марка", "размер", "плотность", "гост", "ту"]
                ):
                    key_chars.append(char_value)
            else:
                char_text = str(char or "").strip()
                if char_text:
                    if ":" in char_text:
                        raw_name, raw_value = char_text.split(":", 1)
                        structured_chars.append({
                            "name": raw_name.strip(),
                            "value": raw_value.strip(),
                        })
                    else:
                        structured_chars.append({
                            "name": "",
                            "value": char_text,
                        })
                    formatted_chars.append(char_text)

        search_query = name
        if key_chars:
            search_query = f"{name} {' '.join(key_chars[:3])}".strip()

        qty = str(pos.get("quantity") or "").strip().lower()
        has_qty = qty and qty not in ["не указано", "нет", "none", "null"]
        is_low_quality = not formatted_chars and not has_qty

        items.append({
            "id": f"{tender_id}:{index}",
            "tender_id": tender_id,
            "tender_title": tender_title,
            "source": source,
            "source_label": source_label,
            "position_name": name,
            "normalized_name": name,
            "quantity": pos.get("quantity"),
            "unit": pos.get("unit"),
            "characteristics": formatted_chars,
            "structured_characteristics": structured_chars,
            "notes": pos.get("notes", ""),
            "search_query": search_query,
            "source_documents": pos.get("source_documents", []),
            "general_requirements_applied": pos.get("general_requirements_applied", False),
            "analog_allowed": pos.get("analog_allowed", False),
            "quality_status": "degraded" if is_low_quality else "good",
        })

    return items


def _build_tender_extraction_group(
    *,
    tender_id: str,
    tender_title: str,
    source: str,
    source_label: str,
    result: Dict[str, Any],
    processing_stats: Dict[str, Any],
) -> Dict[str, Any]:
    positions = result.get("positions") or []
    general_goods_requirements = result.get("general_goods_requirements") or []
    general_requirements = _unique_non_empty([
        req.get("value") if isinstance(req, dict) else req
        for req in general_goods_requirements
    ])
    warnings = _normalize_extraction_warnings(
        result.get("warnings"),
        tender_id=tender_id,
    )

    return {
        "tender_id": tender_id,
        "tender_title": tender_title,
        "source": source,
        "source_label": source_label,
        "items": _build_requirement_items(
            positions=positions,
            tender_id=tender_id,
            tender_title=tender_title,
            source=source,
            source_label=source_label,
        ),
        "warnings": warnings,
        "general_requirements": general_requirements,
        "positions": positions,
        "general_goods_requirements": general_goods_requirements,
        "debug": result.get("debug") or {},
        "extraction_summary": result.get("extraction_summary") or {},
        "processing_stats": processing_stats,
    }


@app.post("/api/ai/extract-tender-requirements")
async def api_extract_tender_requirements(payload: dict = Body(...), db: Session = Depends(get_db)):
    """
    Новое извлечение товарных позиций по всей документации
    по архитектуре блока "ИИ-юрист".
    """
    tender_ids = payload.get("tender_ids") or []
    manual_text = (payload.get("manual_text") or "").strip()

    try:
        # 1. Режим ручного текста
        if manual_text:
            files_data = [{
                "filename": "manual_input.txt",
                "status": "success",
                "error_message": "",
                "extracted_text": manual_text,
                "pages": [{"page_num": 1, "text": manual_text, "tables": []}],
            }]

            result = await run_blocking(
                goods_extraction_service.extract_goods_requirements,
                files_data=files_data,
                tender_id="manual",
                job_id="manual_goods_extraction",
            )
            tender_group = _build_tender_extraction_group(
                tender_id="manual",
                tender_title="Ручной ввод",
                source="manual",
                source_label="Ручной ввод",
                result=result,
                processing_stats={
                    "processed_files": 1,
                    "failed_files": 0,
                    "input_files": 1,
                },
            )
            return {
                **result,
                "tenders": [tender_group],
                "items": tender_group["items"],
                "warnings": tender_group["warnings"],
                "general_requirements": tender_group["general_requirements"],
                "processing_stats": tender_group["processing_stats"],
            }

        # 2. Режим тендеров из CRM
        if not tender_ids:
            raise HTTPException(status_code=400, detail="Не передан ни manual_text, ни tender_ids")

        tender_groups = []
        all_items = []
        all_positions = []
        all_general_goods_requirements = []
        all_warnings = []
        general_requirements_pool: List[str] = []
        processed_files = 0
        failed_files = 0

        for tender_id in tender_ids:
            tender = db.query(TenderModel).filter(TenderModel.id == tender_id).first()
            if not tender:
                all_warnings.extend(_normalize_extraction_warnings(
                    ["Тендер не найден в базе"],
                    tender_id=str(tender_id),
                ))
                continue

            local_path = (getattr(tender, "local_file_path", None) or "").strip()
            extracted_text = (getattr(tender, "extracted_text", None) or "").strip()
            description_text = (getattr(tender, "description", None) or "").strip()

            all_files, has_archives, waited_for_files = await run_blocking(
                wait_for_tender_input_files,
                local_path,
                tender_id,
                DOCUMENTS_ROOT,
                logger=logger,
                max_wait_seconds=15.0,
                poll_interval_seconds=1.0,
            )

            logger.info(
                f"[GOODS_PACKET_START] tender_id={tender_id} selected_files_count={len(all_files)} "
                f"selected_filenames={[os.path.basename(f) for f in all_files]} request_source='CRM' "
                f"has_archives={has_archives} has_local_file_path={bool(local_path)} "
                f"has_extracted_text={bool(extracted_text)} waited_for_files_seconds={waited_for_files:.1f}"
            )

            added_any = False
            packet_files_data = []

            # Process files
            for fpath in all_files:
                filename = os.path.basename(fpath)
                ext = os.path.splitext(fpath)[1].lower()
                size_bytes = os.path.getsize(fpath) if os.path.exists(fpath) else 0

                extracted = await run_blocking(
                    doc_service.extract_document_data,
                    fpath,
                    tender_id=tender_id,
                )
                packet_files_data.append(extracted)

                if extracted.get("status") == "success" and extracted.get("extracted_text"):
                    decision = "included"
                    reason = "Success"
                    added_any = True
                elif ext in ['.zip', '.7z', '.rar']:
                    decision = "archive_skipped"
                    reason = "Archive not supported"
                elif extracted.get("status") == "skipped_empty":
                    decision = "empty"
                    reason = "Empty text"
                else:
                    decision = "skipped"
                    reason = extracted.get("error_message") or "Unknown"

                logger.info(
                    f"[GOODS_PACKET_FILE_DECISION] tender_id={tender_id} filename='{filename}' full_path='{fpath}' "
                    f"extension='{ext}' decision='{decision}' reason='{reason}' priority='N/A' size_bytes={size_bytes}"
                )

            # Fallbacks
            if not added_any and extracted_text:
                packet_files_data.append({
                    "filename": f"{tender_id}_extracted_text.txt",
                    "status": "success",
                    "error_message": "",
                    "extracted_text": extracted_text,
                    "pages": [{"page_num": 1, "text": extracted_text, "tables": []}],
                })
                logger.warning(
                    f"[GOODS_PACKET_FILE_DECISION] tender_id={tender_id} filename='{tender_id}_extracted_text.txt' full_path='N/A' "
                    f"extension='.txt' decision='degraded' reason='Fallback to extracted_text' priority='N/A' size_bytes={len(extracted_text)}"
                )
                added_any = True

            if not added_any and description_text:
                packet_files_data.append({
                    "filename": f"{tender_id}_description.txt",
                    "status": "success",
                    "error_message": "",
                    "extracted_text": description_text,
                    "pages": [{"page_num": 1, "text": description_text, "tables": []}],
                })
                logger.warning(
                    f"[GOODS_PACKET_FILE_DECISION] tender_id={tender_id} filename='{tender_id}_description.txt' full_path='N/A' "
                    f"extension='.txt' decision='degraded' reason='Fallback to description' priority='N/A' size_bytes={len(description_text)}"
                )
                added_any = True

            total_text_chars = sum(len(f.get("extracted_text", "")) for f in packet_files_data if f.get("status") == "success")
            total_pages = sum(len(f.get("pages", [])) for f in packet_files_data if f.get("status") == "success")
            total_tables = sum(sum(len(p.get("tables", [])) for p in f.get("pages", [])) for f in packet_files_data if f.get("status") == "success")
            
            success_files = sum(1 for f in packet_files_data if f.get("status") == "success" and f.get("extracted_text"))
            packet_failed_files = len(packet_files_data) - success_files
            archives_skipped = sum(1 for f in all_files if f.lower().endswith(('.zip', '.7z', '.rar')))
            
            processed_files += success_files
            failed_files += packet_failed_files

            logger.info(
                f"[GOODS_PACKET_READY] tender_id={tender_id} total_input_files={len(all_files)} "
                f"included_files_count={success_files} skipped_files_count={packet_failed_files} archives_skipped_count={archives_skipped} "
                f"success_files={success_files} failed_files={packet_failed_files} total_text_chars={total_text_chars} total_pages={total_pages} "
                f"total_tables={total_tables} critical_degraded={not added_any or (not all_files and bool(extracted_text or description_text))}"
            )

            per_tender_stats = {
                "processed_files": success_files,
                "failed_files": packet_failed_files,
                "input_files": len(all_files),
                "archives_skipped": archives_skipped,
                "total_text_chars": total_text_chars,
                "total_pages": total_pages,
                "total_tables": total_tables,
            }

            if not packet_files_data:
                empty_result = {
                    "positions": [],
                    "general_goods_requirements": [],
                    "warnings": ["Не удалось получить текст ни из одного документа тендера"],
                    "debug": {
                        "documents_count": 0,
                        "included_files": [],
                        "skipped_files": [],
                        "archives_skipped": [],
                        "total_chars": 0,
                    },
                    "extraction_summary": {
                        "positions_count": 0,
                        "duplicates_merged": 0,
                        "ignored_fragments_types": ["CONTRACT_TERMS", "PROCUREMENT_RULES", "PRICE_JUSTIFICATION"],
                        "warnings": ["Не удалось получить текст ни из одного документа тендера"],
                    }
                }
                tender_group = _build_tender_extraction_group(
                    tender_id=tender_id,
                    tender_title=tender.title or tender_id,
                    source="crm",
                    source_label=getattr(tender, "eis_number", None) or tender.id,
                    result=empty_result,
                    processing_stats=per_tender_stats,
                )
            else:
                tender_result = await run_blocking(
                    goods_extraction_service.extract_goods_requirements,
                    files_data=packet_files_data,
                    tender_id=tender_id,
                    job_id=f"crm_goods_extraction_{tender_id}",
                )
                tender_group = _build_tender_extraction_group(
                    tender_id=tender_id,
                    tender_title=tender.title or tender_id,
                    source="crm",
                    source_label=getattr(tender, "eis_number", None) or tender.id,
                    result=tender_result,
                    processing_stats=per_tender_stats,
                )

            tender_groups.append(tender_group)
            all_items.extend(tender_group["items"])
            all_positions.extend(tender_group["positions"])
            all_general_goods_requirements.extend(tender_group["general_goods_requirements"])
            all_warnings.extend(tender_group["warnings"])
            general_requirements_pool.extend(tender_group["general_requirements"])

        if not tender_groups:
            return {
                "tenders": [],
                "items": [],
                "positions": [],
                "general_goods_requirements": [],
                "general_requirements": [],
                "warnings": _normalize_extraction_warnings(
                    ["Не удалось получить текст ни из одного документа тендера"]
                ),
                "debug": {
                    "documents_count": 0,
                    "included_files": [],
                    "skipped_files": [],
                    "archives_skipped": [],
                    "total_chars": 0,
                },
                "extraction_summary": {
                    "positions_count": 0,
                    "duplicates_merged": 0,
                    "ignored_fragments_types": ["CONTRACT_TERMS", "PROCUREMENT_RULES", "PRICE_JUSTIFICATION"],
                    "warnings": ["Не удалось получить текст ни из одного документа тендера"],
                }
            }

        general_requirements = _unique_non_empty(general_requirements_pool)

        return {
            "tenders": tender_groups,
            "items": all_items,
            "positions": all_positions,
            "general_goods_requirements": all_general_goods_requirements,
            "general_requirements": general_requirements,
            "warnings": all_warnings,
            "extraction_summary": {
                "positions_count": len(all_items),
                "duplicates_merged": sum(
                    int((group.get("extraction_summary") or {}).get("duplicates_merged") or 0)
                    for group in tender_groups
                ),
                "ignored_fragments_types": ["CONTRACT_TERMS", "PROCUREMENT_RULES", "PRICE_JUSTIFICATION"],
                "warnings": [warning.get("message") for warning in all_warnings],
            },
            "processing_stats": {
                "processed_files": processed_files,
                "failed_files": failed_files,
                "input_tenders": tender_ids,
                "groups_count": len(tender_groups),
            },
            "debug": {
                "documents_count": sum(
                    int((group.get("processing_stats") or {}).get("input_files") or 0)
                    for group in tender_groups
                ),
            },
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"api_extract_tender_requirements failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Ошибка извлечения ТЗ: {str(e)}")

@app.post("/api/ai/enrich-specs")
async def api_enrich_specs(data: dict = Body(...), _ = Depends(check_ai_service)):
    """Поиск характеристик товара в интернете"""
    logger.info("AI Enrich Specs request.")
    product_name = data.get('product_name', '')
    result = await run_blocking(ai_service.enrich_product_specs, product_name)
    return {"specs": result}

@app.post("/api/ai/match-product")
async def api_match_product(data: dict = Body(...), db: Session = Depends(get_db), _ = Depends(check_ai_service)):
    specs = data.get('specs', '')
    mode = data.get('mode', 'database') # 'database' or 'internet'
    logger.info(f"AI Match Product request. Mode: {mode}, Query len: {len(specs)}")

    if mode == 'internet':
        # Поиск в интернете через Grounding
        result_text = await run_blocking(ai_service.search_products_internet, specs)
        return {"mode": "internet", "text": result_text}
    else:
        # Поиск по базе
        products_db = db.query(ProductModel).limit(50).all()
        catalog = [{"id": str(p.id), "title": p.title, "specs": p.specs} for p in products_db]
        matches = await run_blocking(ai_service.find_product_equivalent, specs, catalog)
        return {"mode": "database", "matches": matches}

@app.post("/api/ai/validate-compliance")
async def api_validate_compliance(data: dict = Body(...), _ = Depends(check_ai_service)):
    """Валидация ТЗ vs Материал (Complex)"""
    logger.info("AI Compliance Validation request.")
    requirements = data.get('requirements', '')
    proposal = data.get('proposal', '[]')
    return await run_blocking(ai_service.compare_requirements_vs_proposal, requirements, proposal)

@app.post("/api/ai/check-compliance")
async def api_check_compliance(data: dict = Body(...), _ = Depends(check_ai_service)):
    """Проверка пакета документов"""
    logger.info("AI Document Package Check request.")
    return await run_blocking(
        ai_service.check_compliance,
        data['title'],
        data['description'],
        data['filenames'],
    )

@app.post("/api/tenders/{tender_id}/upload")
async def upload_tender_file(
    tender_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """Вручную загрузить файл в папку тендера"""
    tender = db.query(TenderModel).filter(TenderModel.id == tender_id).first()
    if not tender:
        raise HTTPException(status_code=404, detail="Тендер не найден")

    tender_dir = os.path.join(DOCUMENTS_ROOT, tender_id)
    os.makedirs(tender_dir, exist_ok=True)

    file_path = os.path.join(tender_dir, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    logger.info(f"Manually uploaded file {file.filename} for tender {tender_id}")
    
    # UNPACK if archive
    try:
        from backend.services.archive_service import archive_service
        if archive_service.is_archive(file_path):
            archive_service.unpack_directory(tender_dir)
            logger.info(f"Unpacked manual upload {file.filename} in {tender_dir}")
    except Exception as e:
        logger.error(f"Error unpacking manual upload {file.filename}: {e}")

    return {
        "status": "success",
        "file": {
            "name": file.filename,
            "size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
            "path": file_path
        }
    }

@app.post("/api/tenders/upload")
async def upload_file(file: UploadFile = File(...), _ = Depends(check_doc_service)):
    logger.info(f"File upload request: {file.filename}")
    
    try:
        file_path = await doc_service.save_file(file)
        # Запускаем OCR или извлечение текста
        text = await run_blocking(doc_service.extract_text, file_path)
        logger.info("File processed successfully.")
        return {"text": text, "path": file_path}
    except Exception as e:
        logger.error(f"Upload Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/dashboard-stats")
async def get_dashboard_stats(db: Session = Depends(get_db)):
    logger.info("Dashboard stats requested.")
    count = db.query(TenderModel).count()
    return {
        "active_tenders": count,
        "margin_val": "₽14.2M",
        "risks_count": 5,
        "contracts_count": 12,
        "chart_data": [{"name": "Пн", "Тендеры": 10, "Выиграно": 2}],
        "tasks": [{"id": "1", "title": "Запустить парсер", "time": "Сейчас", "type": "info"}],
        "is_demo": False
    }

from backend.markdown_parser import add_markdown_to_docx
import zipfile
import io


def _safe_doc_filename(value: str, fallback: str = "export") -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "_", str(value or "").strip())
    cleaned = cleaned.strip(" ._")
    return cleaned[:120] if cleaned else fallback


def _build_parsed_tz_word_document(tenders: List[Dict[str, Any]], selected_item_ids: Optional[set[str]] = None) -> Document:
    selected_item_ids = selected_item_ids or set()
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    doc.add_heading('Спарсенное техническое задание', 0)

    exported_any = False
    for tender_index, tender in enumerate(tenders or [], start=1):
        items = tender.get("items") or []
        if selected_item_ids:
            items = [item for item in items if str(item.get("id") or "") in selected_item_ids]

        if not items and not (tender.get("general_requirements") or []):
            continue

        exported_any = True
        tender_title = str(tender.get("tender_title") or tender.get("title") or f"Тендер {tender_index}")
        tender_id = str(tender.get("tender_id") or tender.get("id") or "").strip()
        source_label = str(tender.get("source_label") or "").strip()

        doc.add_heading(f"{tender_index}. {tender_title}", level=1)
        meta_parts = [part for part in [tender_id, source_label] if part]
        if meta_parts:
            doc.add_paragraph(" | ".join(meta_parts))

        warnings = tender.get("warnings") or []
        if warnings:
            doc.add_heading("Предупреждения", level=2)
            for warning in warnings:
                message = str(warning.get("message") if isinstance(warning, dict) else warning or "").strip()
                if message:
                    doc.add_paragraph(message, style='List Bullet')

        general_requirements = tender.get("general_requirements") or []
        if general_requirements:
            doc.add_heading("Общие требования", level=2)
            for requirement in general_requirements:
                text = str(requirement or "").strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')

        if items:
            doc.add_heading("Позиции", level=2)

        for item_index, item in enumerate(items, start=1):
            position_name = str(item.get("position_name") or f"Позиция {item_index}").strip()
            paragraph = doc.add_paragraph()
            paragraph.add_run(f"{item_index}. {position_name}").bold = True

            qty = str(item.get("quantity") or "").strip()
            unit = str(item.get("unit") or "").strip()
            if qty:
                doc.add_paragraph(f"Количество: {qty} {unit}".strip())

            notes = str(item.get("notes") or "").strip()
            if notes:
                doc.add_paragraph(f"Примечание: {notes}")

            characteristics = item.get("structured_characteristics") or []
            if not characteristics:
                characteristics = [
                    {"name": "", "value": str(value or "").strip()}
                    for value in (item.get("characteristics") or [])
                    if str(value or "").strip()
                ]

            if characteristics:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Характеристика'
                header_cells[1].text = 'Значение'

                for characteristic in characteristics:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(characteristic.get("name") or "Параметр").strip()
                    row_cells[1].text = str(characteristic.get("value") or "").strip()
            else:
                doc.add_paragraph("Характеристики не извлечены.")

    if not exported_any:
        raise HTTPException(status_code=400, detail="Нет данных для экспорта в Word")

    return doc

@app.post("/api/ai/export-risks-word")
async def api_export_risks_word(data: dict = Body(...)):
    """Экспорт результатов анализа рисков в Word .docx (или ZIP для нескольких)"""
    logger.info("Word export started")
    results = data.get('results', [])
    if not results:
        raise HTTPException(status_code=400, detail="No results to export")

    try:
        if len(results) == 1:
            # Single file export
            tender = results[0]
            tid = str(tender.get('id', 'N/A'))
            desc = tender.get('description', 'Нет описания')
            final_report_markdown = tender.get('final_report_markdown') or ""
            summary_notes = tender.get('summary_notes') or ""
            file_statuses = tender.get('file_statuses', [])
            
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)
            
            doc.add_heading('Юридическое заключение по тендеру', 0)
            doc.add_heading(f'Тендер: {tid}', level=1)
            
            p = doc.add_paragraph()
            p.add_run('Описание: ').bold = True
            p.add_run(desc)
            
            if summary_notes:
                doc.add_heading('Краткое резюме:', level=2)
                doc.add_paragraph(summary_notes)
            
            if file_statuses:
                doc.add_heading('Обработанные документы:', level=2)
                for fs in file_statuses:
                    status_text = "Успешно" if fs.get('status') == 'ok' else "Ошибка"
                    doc.add_paragraph(f"{fs.get('filename')} - {status_text}", style='List Bullet')
            
            if not final_report_markdown:
                doc.add_paragraph("Отчет отсутствует.")
            else:
                add_markdown_to_docx(doc, final_report_markdown)
                
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                logger.info(f"Word export finished successfully (single). Path: {tmp.name}")
                return FileResponse(tmp.name, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=f"tender_{tid}_report.docx")
        
        else:
            # Multiple files export (ZIP)
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for tender in results:
                    tid = str(tender.get('id', 'N/A'))
                    desc = tender.get('description', 'Нет описания')
                    final_report_markdown = tender.get('final_report_markdown') or ""
                    summary_notes = tender.get('summary_notes') or ""
                    file_statuses = tender.get('file_statuses', [])
                    
                    doc = Document()
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Arial'
                    font.size = Pt(11)
                    
                    doc.add_heading('Юридическое заключение по тендеру', 0)
                    doc.add_heading(f'Тендер: {tid}', level=1)
                    
                    p = doc.add_paragraph()
                    p.add_run('Описание: ').bold = True
                    p.add_run(desc)
                    
                    if summary_notes:
                        doc.add_heading('Краткое резюме:', level=2)
                        doc.add_paragraph(summary_notes)
                    
                    if file_statuses:
                        doc.add_heading('Обработанные документы:', level=2)
                        for fs in file_statuses:
                            status_text = "Успешно" if fs.get('status') == 'ok' else "Ошибка"
                            doc.add_paragraph(f"{fs.get('filename')} - {status_text}", style='List Bullet')
                    
                    if not final_report_markdown:
                        doc.add_paragraph("Отчет отсутствует.")
                    else:
                        add_markdown_to_docx(doc, final_report_markdown)
                        
                    doc_buffer = io.BytesIO()
                    doc.save(doc_buffer)
                    zip_file.writestr(f"tender_{tid}_report.docx", doc_buffer.getvalue())
            
            zip_buffer.seek(0)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
                tmp.write(zip_buffer.getvalue())
                logger.info(f"Word export finished successfully (zip). Path: {tmp.name}")
                return FileResponse(tmp.name, media_type="application/zip", filename="tenders_reports.zip")

    except Exception as e:
        logger.error(f"Word Export Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/ai/export-parsed-tz-word")
async def api_export_parsed_tz_word(data: dict = Body(...)):
    """Экспорт спарсенного ТЗ по тендеру или выбранным позициям в Word."""
    tenders = data.get("tenders") or []
    selected_item_ids = {
        str(item_id).strip()
        for item_id in (data.get("selected_item_ids") or [])
        if str(item_id).strip()
    }

    if not tenders:
        raise HTTPException(status_code=400, detail="Нет данных ТЗ для экспорта")

    try:
        doc = _build_parsed_tz_word_document(tenders, selected_item_ids)
        first_tender = tenders[0] if tenders else {}
        filename_base = _safe_doc_filename(
            str(first_tender.get("tender_title") or first_tender.get("tender_id") or "parsed_tz"),
            fallback="parsed_tz",
        )

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return FileResponse(
                tmp.name,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{filename_base}_parsed_tz.docx",
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Parsed TZ Word Export Error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

# ─────────────────────────────────────────────────────────────────────────────
# ENDPOINTS: ПОДБОР АНАЛОГОВ
# ─────────────────────────────────────────────────────────────────────────────

@app.post("/api/products/search")
async def search_products_local(
    request: dict,
    db: Session = Depends(get_db),
    _ = Depends(check_analog_service)
):
    """
    Поиск аналогов только в локальной БД.
    """
    query = (request.get("query") or "").strip()
    category = request.get("category")
    requirements = request.get("requirements")
    limit = int(request.get("limit", 10))

    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    reference_profile = await run_blocking(
        analog_service._resolve_reference_profile,
        query,
        requirements or "",
    )
    effective_requirements = analog_service._augment_requirements_with_reference(
        requirements,
        reference_profile,
    )

    results = await run_blocking(
        analog_service.search_local_db,
        query=query,
        category=category,
        requirements=requirements,
        limit=limit,
        reference_profile=reference_profile,
    )
    validation_error = ""
    validation_summary = ""
    if results:
        validated_results, validation_summary, validation_state = await analog_service.ai_validate_candidates(
            query=query,
            requirements=effective_requirements,
            candidates=results,
        )
        if validation_state == "applied":
            results = validated_results
            if not validated_results and validation_summary:
                validation_error = validation_summary
        elif validation_state == "fallback" and validation_summary:
            validation_error = validation_summary
    return {
        "query": query,
        "results": results,
        "total": len(results),
        "ai_error": validation_error,
        "validation_summary": validation_summary,
        "reference_profile": analog_service._public_reference_profile(reference_profile),
    }

@app.post("/api/products/search-ai")
async def search_products_ai(request: dict, background_tasks: BackgroundTasks):
    """Поиск аналогов через Gemini AI с Google Search или в комбинированном режиме."""
    query = str(request.get("query", "")).strip()
    requirements = request.get("requirements")
    max_results = int(request.get("max_results", 5) or 5)
    mode = str(request.get("mode", "both")).strip().lower()
    background_refine = bool(request.get("background_refine", False))

    if not query:
        raise HTTPException(status_code=400, detail="query is required")

    warnings = _build_analog_search_warnings(query)

    if analog_service is None:
        raise HTTPException(status_code=503, detail="AnalogService is not initialized")

    logger.info(f"AI analog search: '{query}' | mode={mode}")

    if background_refine and mode in ["ai", "both"]:
        preview_result = await run_blocking(
            analog_service.build_preview_result,
            query=query,
            requirements=requirements,
            limit=max_results,
        )
        preview_payload = _merge_analog_ai_error(
            _apply_analog_search_mode(preview_result, mode)
        )
        job_id = analog_search_job_service.create_job(
            query=query,
            mode=mode,
            result=preview_payload,
            context={"base_result": preview_result},
        )
        background_tasks.add_task(
            _run_analog_search_refinement_job,
            job_id=job_id,
            query=query,
            requirements=requirements,
            mode=mode,
            max_results=max_results,
        )
        preview_payload["warnings"] = warnings
        preview_payload["background_refine"] = True
        preview_payload["refinement_pending"] = True
        preview_payload["refinement_job"] = analog_search_job_service.get_job(job_id)
        return preview_payload

    result = await analog_service.search_analogs(
        query=query,
        requirements=requirements,
        use_ai=(mode in ["ai", "both"]),
        limit=max_results
    )
    result = _merge_analog_ai_error(_apply_analog_search_mode(result, mode))
    result["warnings"] = warnings
    result["background_refine"] = False
    result["refinement_pending"] = False

    return result


@app.get("/api/products/search-ai/jobs/{job_id}")
async def get_product_search_job_status(job_id: str):
    job = analog_search_job_service.get_job(job_id)
    if not job:
        logger.warning(
            "Analog search refinement job is missing | job_id=%s | returning terminal error payload",
            job_id,
        )
        return analog_search_job_service.build_missing_job_payload(job_id)
    return job

@app.post("/api/products/refresh-catalog")
async def refresh_catalog(background_tasks: BackgroundTasks):
    """Запускает фоновое обновление каталога с gidroizol.ru."""
    async def _do_refresh():
        logger.info("[CATALOG] Starting catalog refresh from gidroizol.ru...")
        try:
            with next(get_db()) as db:
                parse_stats = await parser_service.parse_and_save(db)
            logger.info(
                "[CATALOG] Refresh complete. Parsed %s | inserted=%s | updated=%s | missing_url=%s | unseen_existing=%s.",
                parse_stats.get("parsed", 0),
                parse_stats.get("inserted", 0),
                parse_stats.get("updated", 0),
                parse_stats.get("missing_url", 0),
                parse_stats.get("not_seen_existing", 0),
            )
        except Exception as e:
            logger.error(f"[CATALOG] Refresh failed: {e}")

    background_tasks.add_task(_do_refresh)
    return {"status": "started", "message": "Обновление каталога запущено в фоне"}


@app.post("/api/products")
async def create_product(request: dict, db: Session = Depends(get_db)):
    """Добавить товар в локальную БД вручную."""
    from sqlalchemy import text
    import json as json_lib
    try:
        specs_payload = request.get("specs", {}) or {}
        specs_json = json_lib.dumps(specs_payload, ensure_ascii=False)
        normalized_category = re.sub(r"\s+", " ", str(request.get("category", "")).lower()).strip()
        specs_text = re.sub(
            r"\s+",
            " ",
            " ".join(f"{key} {value}" for key, value in specs_payload.items()),
        ).strip()
        description = request.get("description", "") or ""
        quality_score = 20
        if description:
            quality_score += 10
        if specs_payload:
            quality_score += min(len(specs_payload) * 2, 20)
        result = db.execute(
            text(
                "INSERT INTO products ("
                "vendor, source_url, title, category, normalized_category, searchable_for_analogs, "
                "material_type, price, price_currency, specs, specs_text, url, description, "
                "quality_score, parse_version, is_active"
                ") VALUES ("
                ":vendor, :source_url, :title, :cat, :normalized_category, :searchable, "
                ":mat, :price, :price_currency, :specs, :specs_text, :url, :desc, "
                ":quality_score, :parse_version, 1"
                ")"
            ),
            {
                "vendor": "manual",
                "source_url": request.get("url", ""),
                "title": request.get("title", ""),
                "cat": request.get("category", ""),
                "normalized_category": normalized_category,
                "searchable": 1,
                "mat": request.get("material_type", ""),
                "price": request.get("price"),
                "price_currency": "RUB",
                "specs": specs_json,
                "specs_text": specs_text,
                "url": request.get("url", ""),
                "desc": description,
                "quality_score": quality_score,
                "parse_version": "manual-entry",
            }
        )
        db.commit()
        return {"id": result.lastrowid, "status": "created"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/products/{product_id}")
async def delete_product(product_id: int, db: Session = Depends(get_db)):
    """Удалить товар из БД."""
    from sqlalchemy import text
    db.execute(text("DELETE FROM products WHERE id = :id"), {"id": product_id})
    db.commit()
    return {"status": "deleted", "id": product_id}

# ─────────────────────────────────────────────────────────────────────────────
# КОНЕЦ НОВЫХ ENDPOINTS
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    # log_config=None заставляет uvicorn использовать уже настроенное логирование root-логгера
    uvicorn.run("backend.main:app", host="0.0.0.0", port=8000, reload=True, log_config=None)
