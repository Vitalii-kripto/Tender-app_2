import re
from docx.shared import Pt


def add_markdown_to_docx(doc, md_text):
    md_text = normalize_markdown_tables(md_text or "")
    lines = md_text.replace('\r\n', '\n').split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if '|' in stripped and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.match(r'^\|?[\s\-\|:]+\|?$', next_line) and '|' in next_line:
                table_rows = [stripped, next_line]
                i += 2
                while i < len(lines):
                    candidate = lines[i].strip()
                    if not candidate or '|' not in candidate:
                        break
                    table_rows.append(candidate)
                    i += 1
                _render_table(doc, table_rows)
                continue

        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith('- ') or stripped.startswith('* ') or stripped.startswith('+ '):
            p = doc.add_paragraph(style='List Bullet')
            _parse_inline_formatting(p, stripped[2:].strip())
        elif re.match(r'^\d+[\.\)]\s', stripped):
            p = doc.add_paragraph(style='List Number')
            clean_line = re.sub(r'^\d+[\.\)]\s+', '', stripped).strip()
            _parse_inline_formatting(p, clean_line)
        else:
            p = doc.add_paragraph()
            _parse_inline_formatting(p, stripped)

        i += 1


def normalize_markdown_tables(md_text: str) -> str:
    lines = md_text.replace('\r\n', '\n').split('\n')
    normalized = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if '|' in stripped and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if re.match(r'^\|?[\s\-\|:]+\|?$', next_line) and '|' in next_line:
                block = [stripped, next_line]
                i += 2
                while i < len(lines):
                    candidate = lines[i].strip()
                    if not candidate or '|' not in candidate:
                        break
                    block.append(candidate)
                    i += 1
                normalized.extend(_normalize_table_block(block))
                continue

        normalized.append(line)
        i += 1

    return '\n'.join(normalized)


def _normalize_table_block(table_rows):
    parsed_rows = []
    max_cols = 0

    for row in table_rows:
        if re.match(r'^\|?[\s\-\|:]+\|?$', row) and '|' in row:
            continue
        content = row.strip()
        if content.startswith('|'):
            content = content[1:]
        if content.endswith('|'):
            content = content[:-1]
        cols = [c.strip() for c in content.split('|')]
        parsed_rows.append(cols)
        max_cols = max(max_cols, len(cols))

    if not parsed_rows or max_cols == 0:
        return table_rows

    fixed_rows = []
    header = parsed_rows[0] + [''] * (max_cols - len(parsed_rows[0]))
    fixed_rows.append('| ' + ' | '.join(header) + ' |')
    fixed_rows.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

    for row in parsed_rows[1:]:
        row = row + [''] * (max_cols - len(row))
        fixed_rows.append('| ' + ' | '.join(row) + ' |')

    return fixed_rows


def _render_table(doc, table_rows):
    parsed_rows = []
    for row in table_rows:
        if re.match(r'^\|?[\s\-\|:]+\|?$', row) and '|' in row:
            continue
        content = row.strip()
        if content.startswith('|'):
            content = content[1:]
        if content.endswith('|'):
            content = content[:-1]
        cols = [col.strip() for col in content.split('|')]
        if any(c for c in cols):
            parsed_rows.append(cols)

    if not parsed_rows:
        return

    num_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(parsed_rows):
        row = row + [''] * (num_cols - len(row))
        for j, col_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            _parse_inline_formatting(p, col_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def _parse_inline_formatting(paragraph, text):
    pattern = r'(\*\*\*.*?\*\*\*|___.*?___|\*\*.*?\*\*|__.*?__|(?<!\*)\*.*?\*(?!\*)|\b_.*?_\b)'
    tokens = re.split(pattern, text)

    for token in tokens:
        if not token:
            continue
        if (token.startswith('***') and token.endswith('***')) or (token.startswith('___') and token.endswith('___')):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif (token.startswith('**') and token.endswith('**')) or (token.startswith('__') and token.endswith('__')):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif (token.startswith('*') and token.endswith('*')) or (token.startswith('_') and token.endswith('_')):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)
