import os
os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
os.environ.setdefault("FLAGS_use_mkldnn", "0")
os.environ.setdefault("OMP_NUM_THREADS", "1")

import re
import io
import numpy as np
from pypdf import PdfReader
from fastapi import UploadFile
import aiofiles
import platform
import logging
from typing import Any, Dict, List, Optional

from backend.logger import logger

import pypdfium2 as pdfium
from paddleocr import PaddleOCR

class DocumentService:
    """
    Сервис для работы с документами.
    Поддерживает извлечение текста из PDF и OCR (распознавание сканов) через PaddleOCR.
    """
    UPLOAD_DIR = "uploaded_docs"

    def __init__(self):
        os.makedirs(self.UPLOAD_DIR, exist_ok=True)
        self.text_cache = {}
        self.ocr_engine = None
        self.ocr_init_error = ""
        logger.info("DocumentService initialized in lazy OCR mode.")

    def _ensure_ocr_engine(self):
        if self.ocr_engine is not None:
            return self.ocr_engine
        if self.ocr_init_error:
            raise RuntimeError(self.ocr_init_error)
        try:
            self.ocr_engine = PaddleOCR(lang='ru')
            logger.info("PaddleOCR engine initialized successfully (lazy).")
            return self.ocr_engine
        except Exception as e:
            self.ocr_init_error = f"OCR Configuration Error: {e}"
            logger.error(self.ocr_init_error, exc_info=True)
            raise RuntimeError(self.ocr_init_error)

    def _extract_docx_text_safely(self, file_path: str):
        import docx

        doc = docx.Document(file_path)
        text_parts = []
        tables_data = []

        for para in doc.paragraphs:
            value = (para.text or "").strip()
            if value:
                text_parts.append(value)

        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(cell for cell in row_data):
                    joined = " | ".join(row_data)
                    table_rows.append(joined)
                    text_parts.append(joined)
            if table_rows:
                tables_data.append("\n".join(table_rows))

        for section in doc.sections:
            header = getattr(section, "header", None)
            footer = getattr(section, "footer", None)

            if header:
                for para in header.paragraphs:
                    value = (para.text or "").strip()
                    if value:
                        text_parts.append(value)

            if footer:
                for para in footer.paragraphs:
                    value = (para.text or "").strip()
                    if value:
                        text_parts.append(value)

        full_text = "\n".join(text_parts).strip()
        pages = [{"page_num": 1, "text": full_text}]
        if tables_data:
            pages[0]["tables"] = tables_data

        return full_text, pages

    def _extract_pdf_text_native(self, file_path: str):
        import pdfplumber
        
        pages = []
        text_pages = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    extracted_text = (page.extract_text() or "").strip()
                    extracted_tables = []
                    
                    try:
                        tables = page.extract_tables()
                        for table in tables:
                            table_rows = []
                            for row in table:
                                # Clean row from None and extra spaces
                                clean_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                if any(clean_row):
                                    table_rows.append(" | ".join(clean_row))
                            if table_rows:
                                table_text = "\n".join(table_rows)
                                extracted_tables.append(table_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract tables from PDF page {i+1}: {e}")

                    page_data = {
                        "page_num": i + 1,
                        "text": extracted_text,
                    }
                    if extracted_tables:
                        page_data["tables"] = extracted_tables
                    
                    pages.append(page_data)
                    if extracted_text:
                        text_pages.append(extracted_text)
                    # Add table text to full text too to help classification and AI
                    for t in extracted_tables:
                        text_pages.append(t)
                        
            return "\n".join(text_pages).strip(), pages
        except Exception as e:
            logger.error(f"pdfplumber failed for {file_path}: {e}")
            # Fallback to pypdf (existing logic)
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = []
            text_pages = []
            for i, page in enumerate(reader.pages):
                try:
                    extracted = (page.extract_text() or "").strip()
                except:
                    extracted = ""
                pages.append({"page_num": i + 1, "text": extracted})
                if extracted:
                    text_pages.append(extracted)
            return "\n".join(text_pages).strip(), pages

    def _normalize_extracted_text(self, text: str) -> str:
        text = (text or "").replace("\xa0", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extract_text_from_ocr_result(self, raw_result) -> str:
        page_lines = []

        if raw_result is None:
            return ""

        if isinstance(raw_result, list):
            for item in raw_result:
                if isinstance(item, list):
                    for row in item:
                        if isinstance(row, list) and len(row) >= 2:
                            maybe_text = row[1]
                            if isinstance(maybe_text, (list, tuple)) and maybe_text:
                                value = str(maybe_text[0]).strip()
                                if value:
                                    page_lines.append(value)
                        elif isinstance(row, dict):
                            value = str(row.get("text") or "").strip()
                            if value:
                                page_lines.append(value)
                elif isinstance(item, dict):
                    value = str(item.get("text") or "").strip()
                    if value:
                        page_lines.append(value)

        elif isinstance(raw_result, dict):
            rec_texts = raw_result.get("rec_texts") or raw_result.get("texts") or []
            for value in rec_texts:
                value = str(value).strip()
                if value:
                    page_lines.append(value)

        return "\n".join(page_lines).strip()

    def _ocr_single_image(self, img_array) -> str:
        engine = self._ensure_ocr_engine()
        attempts = [
            ("ocr(img)", lambda: engine.ocr(img_array)),
            ("ocr([img])", lambda: engine.ocr([img_array])),
            ("predict(img)", lambda: engine.predict(img_array)),
            ("predict(input=img)", lambda: engine.predict(input=img_array)),
        ]

        errors = []
        for label, attempt in attempts:
            try:
                raw = attempt()
                text = self._extract_text_from_ocr_result(raw)
                if text:
                    return text
            except Exception as e:
                errors.append(f"{label}: {e}")

        raise RuntimeError(" ; ".join(errors) if errors else "OCR failed with unknown error")

    async def save_file(self, file: UploadFile) -> str:
        """Сохраняет загруженный файл"""
        file_path = os.path.join(self.UPLOAD_DIR, file.filename)
        logger.info(f"Saving file to: {file_path}")
        async with aiofiles.open(file_path, 'wb') as out_file:
            content = await file.read()
            await out_file.write(content)
        return file_path

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """
        Пытается конвертировать .doc в .docx.
        Возвращает путь к новому файлу или оригинальный путь, если не удалось.
        """
        docx_path = doc_path + "x"
        if os.path.exists(docx_path):
            logger.info(f"DOCX version already exists: {docx_path}")
            return docx_path

        # 1. Попытка через win32com (только Windows + Word)
        if platform.system() == "Windows":
            try:
                import win32com.client as win32
                import pythoncom
                # Инициализация COM для текущего потока (важно для FastAPI/async)
                pythoncom.CoInitialize()
                
                word = win32.gencache.EnsureDispatch('Word.Application')
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(doc_path))
                # 16 = wdFormatXMLDocument (.docx)
                doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
                doc.Close()
                word.Quit()
                logger.info(f"Successfully converted {doc_path} to {docx_path} using MS Word.")
                return docx_path
            except Exception as e:
                logger.warning(f"win32com conversion failed (check if Word is installed): {e}")
            finally:
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

        # 2. Попытка через извлечение текста и создание нового .docx (fallback)
        text = self._extract_text_from_doc(doc_path)
        if text and not text.startswith("[ОШИБКА"):
            try:
                import docx
                new_doc = docx.Document()
                new_doc.add_paragraph(f"--- АВТОМАТИЧЕСКАЯ КОНВЕРТАЦИЯ ИЗ .DOC ---\nОригинал: {os.path.basename(doc_path)}\n\n")
                new_doc.add_paragraph(text)
                new_doc.save(docx_path)
                logger.info(f"Created {docx_path} from extracted text of {doc_path}")
                return docx_path
            except Exception as e:
                logger.error(f"Failed to create docx from text: {e}")

        return doc_path

    def _is_html_file(self, file_path: str) -> bool:
        """
        Проверяет, не является ли файл HTML-страницей (бывает при ошибках скачивания с ЕИС).
        """
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(2048)
                # Ищем типичные HTML теги
                chunk_str = chunk.decode('utf-8', errors='ignore').lower()
                if '<html' in chunk_str or '<!doctype html' in chunk_str or '<head' in chunk_str:
                    return True
        except Exception as e:
            logger.warning(f"Error checking if file is HTML: {e}")
        return False

    def extract_document_data(self, file_path: str, use_cache: bool = True, tender_id: str = "unknown") -> Dict[str, Any]:
        """
        Извлечение структурированных данных из документа (страницы, листы, статус).
        """
        filename = os.path.basename(file_path)
        ext = os.path.splitext(file_path)[1].lower()
        
        logger.info(f"[DOC_EXTRACT_START] tender_id={tender_id} filename='{filename}' extension='{ext}' file_path='{file_path}'")

        result = {
            "filename": filename,
            "file_path": file_path,
            "file_extension": ext,
            "file_size": 0,
            "status": "failed_unknown",
            "extract_method": "none",
            "text_length": 0,
            "extracted_text": "",
            "error_message": "",
            "pages": []
        }

        if not os.path.exists(file_path):
            result["status"] = "skipped_not_found"
            result["error_message"] = "Файл не найден"
            logger.info(f"[DOC_EXTRACT_RESULT] tender_id={tender_id} filename='{filename}' extension='{ext}' status='{result['status']}' extract_method='none' text_length=0 pages_count=0 tables_count=0 quality_flags='' error_message='{result['error_message']}'")
            return result
            
        file_size = os.path.getsize(file_path)
        result["file_size"] = file_size

        if file_size == 0:
            result["status"] = "skipped_empty"
            result["error_message"] = "Файл пуст"
            logger.info(f"[DOC_EXTRACT_RESULT] tender_id={tender_id} filename='{filename}' extension='{ext}' status='{result['status']}' extract_method='none' text_length=0 pages_count=0 tables_count=0 quality_flags='' error_message='{result['error_message']}'")
            return result

        if ext in ['.zip', '.7z', '.rar']:
            result["status"] = "skipped_unsupported_type"
            result["error_message"] = "Archive not supported for direct analysis"
            logger.warning(f"[DOC_ARCHIVE_UNSUPPORTED] tender_id={tender_id} filename='{filename}' extension='{ext}' reason=archive_not_supported_for_direct_analysis")
            logger.info(f"[DOC_EXTRACT_RESULT] tender_id={tender_id} filename='{filename}' extension='{ext}' status='{result['status']}' extract_method='none' text_length=0 pages_count=0 tables_count=0 quality_flags='' error_message='{result['error_message']}'")
            return result

        if self._is_html_file(file_path):
            result["status"] = "skipped_invalid_file"
            result["error_message"] = "Файл является HTML-страницей (ошибка скачивания)"
            logger.info(f"[DOC_EXTRACT_RESULT] tender_id={tender_id} filename='{filename}' extension='{ext}' status='{result['status']}' extract_method='none' text_length=0 pages_count=0 tables_count=0 quality_flags='' error_message='{result['error_message']}'")
            return result

        mtime = os.path.getmtime(file_path)
        cache_key = (file_path, mtime)
        
        if use_cache and cache_key in self.text_cache:
            logger.info(f"Returning cached text for {file_path}")
            cached = self.text_cache[cache_key]
            # Update result with cached data
            result.update({
                "extracted_text": cached.get("text", ""),
                "text_length": len(cached.get("text", "")),
                "status": cached.get("status", "success"),
                "error_message": cached.get("error_message", ""),
                "pages": cached.get("pages", []),
                "extract_method": cached.get("extract_method", "cache")
            })
            return result

        logger.info(f"--- [START STRUCTURED EXTRACTION] ---")
        logger.info(f"File path: {file_path}")
        logger.info(f"Extension: {ext}")
        
        try:
            if ext == '.doc':
                result["extract_method"] = "antiword/striprtf"
                docx_path = file_path + "x"
                if os.path.exists(docx_path):
                    file_path = docx_path
                    ext = '.docx'
                else:
                    new_path = self._convert_doc_to_docx(file_path)
                    if new_path.endswith('.docx'):
                        file_path = new_path
                        ext = '.docx'
                    else:
                        full_text = self._extract_text_from_doc(file_path)
                        result["extracted_text"] = full_text
                        result["pages"] = [{"page_num": 1, "text": full_text}]
                        if not full_text.strip() or full_text.startswith("[ОШИБКА"):
                            result["status"] = "failed_text_extraction"
                            result["error_message"] = full_text
                        else:
                            result["status"] = "success"
                        
                        result["text_length"] = len(result["extracted_text"])
                        return result

            if ext == '.docx':
                result["extract_method"] = "python-docx"
                full_text, pages = self._extract_docx_text_safely(file_path)
                result["extracted_text"] = self._normalize_extracted_text(full_text)
                result["pages"] = pages
                result["status"] = "success"
            
            elif ext == '.xlsx':
                result["extract_method"] = "openpyxl"
                import openpyxl
                wb = openpyxl.load_workbook(file_path, data_only=True)
                sheets_text = []
                for sheet in wb.worksheets:
                    sheet_data = []
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip().replace("|", "").strip():
                            sheet_data.append(row_text)
                    if sheet_data:
                        sheet_text = f"=== ЛИСТ: {sheet.title} ===\n" + "\n".join(sheet_data)
                        sheets_text.append(sheet_text)
                        result["pages"].append({"page_num": sheet.title, "text": sheet_text, "is_sheet": True, "tables": ["\n".join(sheet_data)]})
                result["extracted_text"] = "\n\n".join(sheets_text)
                result["status"] = "success"
            
            elif ext == '.xls':
                result["extract_method"] = "xlrd"
                import xlrd
                wb = xlrd.open_workbook(file_path)
                sheets_text = []
                for sheet in wb.sheets():
                    sheet_data = []
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip().replace("|", "").strip():
                            sheet_data.append(row_text)
                    if sheet_data:
                        sheet_text = f"=== ЛИСТ: {sheet.name} ===\n" + "\n".join(sheet_data)
                        sheets_text.append(sheet_text)
                        result["pages"].append({"page_num": sheet.name, "text": sheet_text, "is_sheet": True, "tables": ["\n".join(sheet_data)]})
                result["extracted_text"] = "\n\n".join(sheets_text)
                result["status"] = "success"
            
            elif ext == '.pdf':
                result["extract_method"] = "pypdf"
                native_text, native_pages = self._extract_pdf_text_native(file_path)
                native_text = self._normalize_extracted_text(native_text)
                result["pages"] = native_pages

                if native_text and self._is_text_quality_good(native_text):
                    result["extracted_text"] = native_text
                    result["status"] = "success"
                else:
                    logger.info("PDF text quality is POOR. Triggering OCR fallback...")
                    result["extract_method"] = "paddleocr"
                    try:
                        ocr_pages = self._perform_ocr(file_path)
                        ocr_text = self._normalize_extracted_text(
                            "\n\n".join([f"--- Page {p['page_num']} ---\n{p['text']}" for p in ocr_pages])
                        )

                        if ocr_text:
                            result["pages"] = ocr_pages
                            result["extracted_text"] = ocr_text
                            result["ocr_used"] = True
                            result["status"] = "success"
                        elif native_text:
                            result["pages"] = native_pages
                            result["extracted_text"] = native_text
                            result["ocr_used"] = True
                            result["status"] = "success"
                            result["error_message"] = "OCR не дал результата, использован нативно извлеченный текст PDF"
                        else:
                            result["pages"] = native_pages
                            result["extracted_text"] = ""
                            result["ocr_used"] = True
                            result["status"] = "failed_ocr"
                            result["error_message"] = "OCR не дал результата, а нативный текст PDF пуст"
                    except Exception as e:
                        logger.error(f"OCR failed: {e}", exc_info=True)
                        result["ocr_used"] = True

                        if native_text:
                            result["pages"] = native_pages
                            result["extracted_text"] = native_text
                            result["status"] = "success"
                            result["error_message"] = f"OCR failed, использован нативный текст PDF: {str(e)}"
                        else:
                            result["pages"] = native_pages
                            result["extracted_text"] = ""
                            result["status"] = "failed_ocr"
                            result["error_message"] = f"OCR failed: {str(e)}"

                filename_lower = filename.lower()
                result["critical_for_analysis"] = any(x in filename_lower for x in ["нмцк", "обоснование", "смет"]) and not result.get("extracted_text")

            else:
                result["status"] = "skipped_unsupported_type"
                result["error_message"] = f"Unsupported file extension: {ext}"

            if result["status"] == "success" and not result["extracted_text"].strip():
                result["status"] = "skipped_empty"
                result["error_message"] = "Файл пуст или текст не извлечен"

        except Exception as e:
            logger.error(f"Extraction error for {file_path}: {e}", exc_info=True)
            result["status"] = "failed_text_extraction"
            result["error_message"] = str(e)

        result["text_length"] = len(result.get("extracted_text", ""))

        if result.get("status") == "success":
            # For cache compatibility, we store it in the old format too
            cache_data = result.copy()
            cache_data["text"] = result["extracted_text"]
            self.text_cache[cache_key] = cache_data

        # Final DOC_EXTRACT_RESULT log
        tables_total = sum(len(p.get("tables", [])) for p in result.get("pages", []))
        pages_count = len(result.get("pages", []))
        logger.info(
            f"[DOC_EXTRACT_RESULT] tender_id={tender_id} filename='{filename}' extension='{ext}' status='{result['status']}' "
            f"extract_method='{result['extract_method']}' text_length={result['text_length']} pages_count={pages_count} "
            f"tables_count={tables_total} quality_flags='' error_message='{result['error_message']}'"
        )
        
        # DOC_PAGE_STATS logs
        for p in result.get("pages", []):
            pnum = p.get("page_num", 0)
            plen = len(p.get("text", ""))
            ptab = len(p.get("tables", []))
            ocr_used = result.get("ocr_used", False)
            # degraded could mean skipped empty due to low quality but let's default to False
            degraded = False
            logger.debug(
                f"[DOC_PAGE_STATS] tender_id={tender_id} filename='{filename}' page_num={pnum} "
                f"text_length={plen} tables_count={ptab} ocr_used={ocr_used} degraded={degraded} quality_score='N/A'"
            )

        return result

    def extract_text(self, file_path: str) -> str:
        """
        Обертка для обратной совместимости.
        Возвращает именно extracted_text, а не несуществующий ключ text.
        """
        data = self.extract_document_data(file_path)
        if data["status"] not in {"success"}:
            logger.warning("Extraction issue: %s - %s", data["status"], data["error_message"])
        return data.get("extracted_text", "")

    def _is_text_quality_good(self, text: str) -> bool:
        """
        Проверяет качество извлеченного текста.
        Возвращает True, если текст качественный, и False, если требуется OCR.
        Ужесточенные критерии для PDF.
        """
        if not text or len(text.strip()) < 150: # Увеличили порог минимальной длины
            return False

        total_chars = len(text)
        
        # 1. Доля мусорных символов
        clean_pattern = r'[a-zA-Zа-яА-ЯёЁ0-9\s\.,!?;:()""\'\'\-\+=\[\]/\\<>@#\$%\^&\*«»№]'
        clean_chars_count = len(re.findall(clean_pattern, text))
        garbage_ratio = 1 - (clean_chars_count / total_chars) if total_chars > 0 else 1
        
        # 2. Доля нормальных слов на русском
        words = text.split()
        if not words:
            return False
            
        russian_words = [w for w in words if re.search(r'[а-яА-ЯёЁ]{3,}', w)]
        russian_words_ratio = len(russian_words) / len(words) if words else 0
        
        # 3. Наличие длинных испорченных строк
        max_word_len = max(len(w) for w in words)
        avg_word_len = sum(len(w) for w in words) / len(words)
        
        # 4. Доля нечитаемых фрагментов
        unreadable_chars = len(re.findall(r'[\?\x00-\x08\x0b\x0c\x0e-\x1f]', text))
        unreadable_ratio = unreadable_chars / total_chars if total_chars > 0 else 0

        logger.info(f"PDF Quality Metrics: Garbage={garbage_ratio:.2f}, RusWords={russian_words_ratio:.2f}, MaxWord={max_word_len}, AvgWord={avg_word_len:.1f}, Unreadable={unreadable_ratio:.2f}")

        # Ужесточенные пороговые значения:
        # - Мусора > 10% (было 15%)
        # - Русских слов < 30% (было 20%)
        # - Слишком длинные "слова" (> 80 символов, было 100)
        # - Слишком много нечитаемых знаков (> 3%, было 5%)
        
        if garbage_ratio > 0.10:
            return False
        if russian_words_ratio < 0.30:
            # Если это не технический документ на английском
            english_words = [w for w in words if re.search(r'[a-zA-Z]{3,}', w)]
            english_words_ratio = len(english_words) / len(words)
            if english_words_ratio < 0.40: # И не английский тоже
                return False
        if max_word_len > 80 or avg_word_len > 18:
            return False
        if unreadable_ratio > 0.03:
            return False

        return True

    def _perform_ocr(self, file_path: str) -> List[Dict[str, Any]]:
        import time
        start_time = time.time()
        logger.info(f"Starting OCR for {file_path} using pypdfium2 + PaddleOCR")

        pdf = pdfium.PdfDocument(file_path)
        pages_to_process = min(len(pdf), 100)
        ocr_pages = []

        try:
            for i in range(pages_to_process):
                logger.info(f"Processing page {i + 1}/{pages_to_process}...")
                page = pdf[i]
                bitmap = page.render(scale=3)
                pil_image = bitmap.to_pil()
                img_array = np.array(pil_image)

                page_text = self._ocr_single_image(img_array)
                page_text = self._normalize_extracted_text(page_text)
                ocr_pages.append({"page_num": i + 1, "text": page_text})
        finally:
            pdf.close()

        full_ocr_text = "\n\n".join([p["text"] for p in ocr_pages]).strip()
        end_time = time.time()

        if full_ocr_text:
            logger.info(
                f"OCR successful. Extracted {len(full_ocr_text)} characters from "
                f"{pages_to_process} pages in {end_time - start_time:.2f}s."
            )
            return ocr_pages

        logger.warning(f"OCR ran but found no text in {end_time - start_time:.2f}s.")
        return [{"page_num": 1, "text": ""}]

    def _extract_text_from_doc(self, file_path: str) -> str:
        """
        Извлечение текста из старого формата .doc.
        Использует striprtf (если это RTF) или системную команду antiword.
        """
        # 1. Попытка через striprtf (некоторые .doc - это на самом деле RTF)
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                if "{\\rtf" in content:
                    text = rtf_to_text(content)
                    if text.strip():
                        logger.info(f"Striprtf extracted {len(text)} characters from .doc (RTF)")
                        return text
        except Exception as e:
            logger.warning(f"Striprtf failed for .doc: {e}")

        # 2. Попытка через системную команду antiword
        import shutil
        import subprocess
        
        antiword_cmd = shutil.which('antiword')
        if antiword_cmd:
            try:
                result = subprocess.run([antiword_cmd, file_path], capture_output=True, text=True, errors='ignore')
                if result.returncode == 0 and result.stdout.strip():
                    logger.info(f"Antiword extracted {len(result.stdout)} characters from .doc")
                    return result.stdout
            except Exception as e:
                logger.warning(f"Antiword execution failed: {e}")
        else:
            logger.warning("Antiword executable not found in PATH.")

        # 3. Последняя попытка: чтение как простого текста (иногда помогает для очень старых или поврежденных файлов)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                raw_content = f.read()
                # Пытаемся найти хоть какой-то осмысленный текст среди бинарных данных
                import re
                clean_text = re.sub(r'[^\x20-\x7E\u0400-\u04FF\n\t]', ' ', raw_content)
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                if len(clean_text) > 100:
                    logger.info("Extracted partial text from .doc using raw fallback.")
                    return f"[ВНИМАНИЕ: Текст извлечен частично]\n\n{clean_text}"
        except:
            pass

        msg = f"Не удалось прочитать файл {os.path.basename(file_path)}."
        if platform.system() == "Windows":
            msg += "\n\nДЛЯ ИСПРАВЛЕНИЯ:\n1. Установите утилиту Antiword и добавьте её в PATH.\n2. ИЛИ (проще) пересохраните файл в формате .docx."
        else:
            msg += "\n\nУстановите пакет antiword (sudo apt install antiword)."
            
        return f"[ОШИБКА ФОРМАТА] {msg}"
